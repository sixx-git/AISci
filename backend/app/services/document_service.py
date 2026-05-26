import os
import uuid
from datetime import datetime
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from app.schemas.documents import DocumentResponse
from app.models import Document
from app.core.config import get_settings

settings = get_settings()


class DocumentService:
    def __init__(self, db: Session):
        self.db = db
        self.storage_path = settings.UPLOAD_DIR
        os.makedirs(self.storage_path, exist_ok=True)
    
    async def upload_document(self, file: UploadFile) -> DocumentResponse:
        # 验证文件扩展名
        file_extension = os.path.splitext(file.filename)[1].lower()
        if not file_extension:
            file_extension = ".txt"
        
        ext = file_extension[1:] if file_extension else ""
        if ext not in settings.allowed_extensions_list:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型。支持的类型: {', '.join(settings.allowed_extensions_list)}"
            )
        
        document_id = str(uuid.uuid4())
        save_filename = f"{document_id}{file_extension}"
        save_path = os.path.join(self.storage_path, save_filename)
        
        with open(save_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        content_text = await self._extract_text(save_path, file.filename)
        
        document = Document(
            id=document_id,
            filename=file.filename,
            file_path=save_path,
            file_type=ext,
            content=content_text,
            status="processed"
        )
        self.db.add(document)
        self.db.commit()
        
        # 向量索引由 /api/v1/vector-search/build 统一构建，不在此处触发
        
        return DocumentResponse(
            success=True,
            document_id=document_id,
            filename=file.filename,
            upload_time=datetime.now(),
            status="processed"
        )
    
    async def list_documents(self):
        documents = self.db.query(Document).all()
        return [
            {
                "id": doc.id,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "status": doc.status,
                "created_at": doc.created_at
            }
            for doc in documents
        ]
    
    async def _extract_text(self, file_path: str, filename: str) -> str:
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                    return f.read()
        elif file_extension == ".md":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif file_extension == ".csv":
            import csv
            text_content = []
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                for row in reader:
                    text_content.append(" ".join(row))
            return "\n".join(text_content)
        elif file_extension == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
                return "\n".join(text_parts)
            except ImportError:
                return f"PDF 文件 {filename} (需安装 PyPDF2 以解析内容)"
        elif file_extension == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text_parts = []
                for paragraph in doc.paragraphs:
                    text_parts.append(paragraph.text)
                return "\n".join(text_parts)
            except ImportError:
                return f"Word 文件 {filename} (需安装 python-docx 以解析内容)"
        else:
            return f"文件 {filename} (文本解析待实现)"
