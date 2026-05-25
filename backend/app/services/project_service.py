"""
项目服务层
"""
import os
import uuid
from typing import List, Optional, Tuple
from datetime import datetime
from sqlalchemy.orm import Session
from sqlalchemy import and_, or_

from app.core.config import get_settings
from app.models.project import Project, ProjectStatus, Document
from app.schemas.project import ProjectCreate, ProjectUpdate, ProjectQuery

settings = get_settings()


class ProjectService:
    """项目服务"""
    
    def __init__(self, db: Session):
        self.db = db
    
    def create_project(self, data: ProjectCreate) -> Project:
        """创建项目"""
        project_id = str(uuid.uuid4())
        project = Project(
            id=project_id,
            name=data.name,
            description=data.description,
            keywords=data.keywords,
            status=ProjectStatus.DRAFT,
            created_by=data.created_by
        )
        
        self.db.add(project)
        self.db.commit()
        self.db.refresh(project)
        
        return project
    
    def get_project(self, project_id: str) -> Optional[Project]:
        """获取项目详情"""
        return self.db.query(Project).filter(Project.id == project_id).first()
    
    def list_projects(
        self,
        query: ProjectQuery
    ) -> Tuple[List[Project], int]:
        """项目列表"""
        q = self.db.query(Project)
        
        # 状态筛选
        if query.status:
            q = q.filter(Project.status == query.status)
        
        # 关键词搜索
        if query.keyword:
            keyword = f"%{query.keyword}%"
            q = q.filter(
                or_(
                    Project.name.like(keyword),
                    Project.description.like(keyword),
                    Project.keywords.like(keyword)
                )
            )
        
        # 排序
        q = q.order_by(Project.created_at.desc())
        
        # 总数
        total = q.count()
        
        # 分页
        offset = (query.page - 1) * query.page_size
        projects = q.offset(offset).limit(query.page_size).all()
        
        return projects, total
    
    def update_project(self, project_id: str, data: ProjectUpdate) -> Optional[Project]:
        """更新项目"""
        project = self.get_project(project_id)
        if not project:
            return None
        
        # 更新字段
        update_data = data.model_dump(exclude_unset=True)
        for field, value in update_data.items():
            setattr(project, field, value)
        
        self.db.commit()
        self.db.refresh(project)
        
        return project
    
    def delete_project(self, project_id: str) -> bool:
        """删除项目"""
        project = self.get_project(project_id)
        if not project:
            return False
        
        self.db.delete(project)
        self.db.commit()
        
        return True


class DocumentService:
    """文档服务"""
    
    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = settings.UPLOAD_DIR
        
        # 确保上传目录存在
        os.makedirs(self.upload_dir, exist_ok=True)
    
    def save_document(
        self,
        filename: str,
        file_content: bytes,
        project_id: Optional[str] = None
    ) -> Document:
        """保存文档"""
        doc_id = str(uuid.uuid4())
        file_extension = os.path.splitext(filename)[1].lower()
        save_filename = f"{doc_id}{file_extension}"
        save_path = os.path.join(self.upload_dir, save_filename)
        
        # 保存文件
        with open(save_path, 'wb') as f:
            f.write(file_content)
        
        # 创建文档记录
        doc = Document(
            id=doc_id,
            project_id=project_id,
            filename=filename,
            file_path=save_path,
            file_type=file_extension[1:] if file_extension else "unknown",
            file_size=len(file_content),
            status="pending"
        )
        
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        
        return doc
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """获取文档"""
        return self.db.query(Document).filter(Document.id == doc_id).first()
    
    def list_documents(
        self,
        project_id: Optional[str] = None,
        page: int = 1,
        page_size: int = 20
    ) -> Tuple[List[Document], int]:
        """文档列表"""
        q = self.db.query(Document)
        
        if project_id:
            q = q.filter(Document.project_id == project_id)
        
        q = q.order_by(Document.created_at.desc())
        
        total = q.count()
        offset = (page - 1) * page_size
        documents = q.offset(offset).limit(page_size).all()
        
        return documents, total
    
    def update_document_status(
        self,
        doc_id: str,
        status: str,
        error_message: Optional[str] = None,
        content: Optional[str] = None,
        summary: Optional[str] = None
    ) -> Optional[Document]:
        """更新文档状态"""
        doc = self.get_document(doc_id)
        if not doc:
            return None
        
        doc.status = status
        if error_message:
            doc.error_message = error_message
        if content:
            doc.content = content
        if summary:
            doc.summary = summary
        
        self.db.commit()
        self.db.refresh(doc)
        
        return doc
    
    def delete_document(self, doc_id: str) -> bool:
        """删除文档"""
        doc = self.get_document(doc_id)
        if not doc:
            return False
        
        # 删除文件
        if os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except:
                pass
        
        self.db.delete(doc)
        self.db.commit()
        
        return True
    
    def extract_text_from_file(self, file_path: str, filename: str) -> str:
        """从文件中提取文本"""
        import os
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return self._extract_docx_text(file_path)
        elif file_extension in ['.txt', '.md']:
            return self._extract_text_file(file_path)
        else:
            return f"不支持的文件格式: {file_extension}"
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """提取 PDF 文本"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text() or ""
                text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            return f"PDF 解析失败: {str(e)}"
    
    def _extract_docx_text(self, file_path: str) -> str:
        """提取 DOCX 文本"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            return "\n".join(text_parts)
        except Exception as e:
            return f"DOCX 解析失败: {str(e)}"
    
    def _extract_text_file(self, file_path: str) -> str:
        """提取文本文件内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception as e:
                return f"文本文件读取失败: {str(e)}"
