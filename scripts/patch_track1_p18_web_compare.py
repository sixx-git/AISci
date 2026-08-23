# -*- coding: utf-8 -*-
"""Fill P18 with the completed web-Qwen 5-round comparison (sjtu_q_087)."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx import Document
from retune_track1_1a_theme import HZ, TPL, set_cell, set_fill  # noqa: E402

P029 = (
    "相较直接使用大模型回答，本作品最主要的不同：用同一套五轮提示直接问网页 Qwen，"
    "也能写出缺口、反对证据和计划，但文献无法核验，不可执行计划不会被记成第一轮失败。"
    "本作品用三大创新把四步中间产物约束并冻结：创新一只允许引用 Fact 白名单并保留反对证据；"
    "创新二拦截不可证伪或无可执行步骤的方案；大家长审核各阶段输入输出，启动或终止证据链迭代，"
    "上一轮已审核结果不得被下一轮覆盖。证明的是四步能否被这套手段做出来，不是「同一道题答得更像标准答案」。"
)

P143 = (
    "科学逻辑方面，本作品实际改善了什么：同一套五轮提示下，网页 Qwen 也能写出缺口、反对证据和计划；"
    "本作品改善的是把这四步做成可回溯、可门禁、可冻结的中间产物，并把第一轮不可执行计划原样留给第二轮，"
    "而不是把万人 RCT 直接当成优先方案。"
)

TABLE20_R1 = [
    "同一题 sjtu_q_087：网页直接问 Qwen（五轮提示、不检索、不走 Pipeline）vs 本作品第一轮冻结件（run_id=c6e7fbfb-645a-461e-8e57-5ceb515bd86a）",
    "同一官方题干与背景；同一类 Qwen 基座；同一套五轮提问（问题理解→缺口→三条含反对证据的假设→筛选→计划）。直接问答不给 Fact 白名单、不检索、不跑门禁。",
    "核对表：①有无可处理缺口；②三条候选是否含反对证据；③有无无法核验的文献；④筛选理由能否复核；⑤计划当前能否执行（网页入选 H-01 的 3–4 万例 RCT、本作品入选 H-03 的万例前瞻均照实写）；⑥第一轮能否原样对照。",
    "缺口围绕合成数据分布偏移与人机分工标准；三条假设均绑定 fact_id 并含反对证据；入选 H-03（三维责任矩阵），Accept 7.04；验证写成万例前瞻，计划不可执行，已原样冻结。",
    "同一套五轮提示下，网页 Qwen 写出 4 条缺口、H-01/H-02/H-03（均含反对证据）、筛选理由和计划，并非只会给是/否。但文献写成「多项研究」而无完整题名、作者、年份，无法核验；入选 H-01（肺结节 LDCT 硬终点 RCT，合计约 3–4 万人、随访 3 年），当前同样不可执行；无 fact_id、无冻结 run。",
    "同提示下差别不在「会不会写四步」，而在四步有没有被白名单、可执行性门禁和第一轮冻结约束住。网页把不可执行的万人 RCT 写成优先方案；本作品把不可执行记成第一轮问题并进入第二轮。代价是步骤与 Token 更多。",
]

HZ_COMPARE = (
    "对照2：对象：同一题 sjtu_q_087，网页直接问 Qwen（五轮提示、不检索、不走 Pipeline）"
    "vs 本作品第一轮冻结件（run_id=c6e7fbfb-645a-461e-8e57-5ceb515bd86a）；"
    "同条件：同一官方题干与背景、同一类 Qwen 基座、同一套五轮提问；"
    "方法：核对表（缺口 / 反对证据 / 无法核验文献 / 筛选理由 / 计划当前能否执行 / 第一轮能否对照）；"
    "本作品：缺口围绕合成数据分布偏移与人机分工标准；三条假设绑定 fact_id 并含反对证据；"
    "入选 H-03（三维责任矩阵），Accept 7.04；验证写成万例前瞻，计划不可执行并已冻结；"
    "对照：网页 Qwen 也能写出 4 条缺口、三条含反对证据的假设、筛选理由和计划，并非只会给是/否，"
    "但文献写成「多项研究」而无完整题名作者年份，无法核验；入选 H-01（肺结节 LDCT 硬终点 RCT，"
    "合计约 3–4 万人、随访 3 年），当前同样不可执行；无 fact_id、无冻结 run；"
    "结论：同提示下差别不在会不会写四步，而在四步有没有被白名单、门禁和第一轮冻结约束；代价是步骤与 Token 更多。"
)

HZ_NOTE = (
    "说明：对照已按「网页对比.docx」完成五轮，无需再补对话。"
    "未做消融；未做与 The AI Scientist 的对照。P18 仅保留这一行同条件比较。"
)

HZ_DIFF = (
    "与直接使用大模型回答的不同："
    "用同一套五轮提示直接问网页 Qwen，也能写出缺口、反对证据和计划，但文献无法核验，"
    "不可执行计划不会被记成第一轮失败。"
    "本作品用三大创新把四步中间产物约束并冻结：创新一只允许引用 Fact 白名单并保留反对证据；"
    "创新二拦截不可证伪或无可执行步骤的方案；大家长审核各阶段输入输出，启动或终止证据链迭代，"
    "上一轮已审核结果不得被下一轮覆盖。"
)


def main() -> None:
    doc = Document(str(TPL))
    set_fill(doc.paragraphs[29], P029)
    set_fill(doc.paragraphs[143], P143)
    t20 = doc.tables[20]
    for c, val in enumerate(TABLE20_R1):
        set_cell(t20.rows[1].cells[c], val)
    doc.save(str(TPL))
    print("template P18 saved")

    hz = Document(str(HZ))
    for p in hz.paragraphs:
        raw = "".join((r.text or "") for r in p.runs)
        if raw.startswith("对照2："):
            set_fill(p, HZ_COMPARE)
        elif raw.startswith("说明：消融实验未做过") or (
            raw.startswith("说明：") and "P18" in raw and "对照" in raw
        ):
            set_fill(p, HZ_NOTE)
        elif raw.startswith("与直接使用大模型回答的不同："):
            set_fill(p, HZ_DIFF)
        elif "科学逻辑方面，本作品实际改善了什么：" in raw[:80]:
            head = raw.split("科学逻辑方面，本作品实际改善了什么：", 1)[0]
            set_fill(p, head + P143)
    hz.save(str(HZ))
    print("huizong P18 saved")


if __name__ == "__main__":
    main()
