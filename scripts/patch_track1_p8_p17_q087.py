# -*- coding: utf-8 -*-
"""Patch P8/P10/P13-P17 in track-1 template and 汇总 draft to sjtu_q_087."""
from __future__ import annotations

from pathlib import Path

from docx import Document

TPL = Path(r"d:\Workplace\AISci\output\提交\模板\赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx")
HZ = Path(r"d:\Workplace\AISci\output\提交\赛道一模板填写稿_汇总.docx")


def set_para(p, text: str) -> None:
    text = text.replace("\r\n", "\n")
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def set_cell(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para(paras[0], text)
    for extra in paras[1:]:
        set_para(extra, "")


P8 = (
    "以官方 125 题中「人工智能能否取代医生？」（sjtu_q_087，信息科学）为例。"
    "项目内展开为：复杂临床场景下医疗人工智能辅助诊断系统的决策可信度边界及人机协同责任分配机制研究。"
    "\n研究对象：内——医疗AI模型的不确定性量化、可解释性与多模态融合；外——诊疗规范、侵权责任与医患信任；边界——影像/病理等辅助诊断，而非自主治疗。"
    "\n关键变量：预测不确定性、病例复杂度、医生经验水平、责任分配权重。"
    "\n已有认识：AI在部分医疗任务上可超过人类，但实施障碍将长期阻碍大规模替代医生；部署须处理伦理与责任（如 fact_003、fact_004，均绑定 source_chunk_id）。"
    "\n知识缺口（第一轮）：①用合成数据补稀缺高危场景时，分布偏移会改变可信度边界，且缺少人机协同责任评估标准；②高置信错误建议下，医生难以判断可信度边界，动态责任分配难以落地。"
    "\n约束：必须符合现行《执业医师法》及医疗侵权责任框架；AI须具备不确定性量化或可解释输出；研究场景限定为辅助诊断而非自主治疗干预。"
    "\n该解析结果直接约束后续假设：不得把「完全取代医生」写成可立即临床部署的结论。"
)

P10 = (
    "筛选结果由“门禁链 + 锦标赛/集成评审”共同形成：候选假设依次经过领域相关性过滤、知识库对齐、引用校验、可证伪性过滤、去重与集成门禁；通过者排序并附选择理由。"
    "以 sjtu_q_087 第一轮真实结果为例（三条假设 evidence_level 均为 high，不是低证据淘汰）："
    "入选 H-03「不确定性×病例复杂度×医生经验的三维动态责任分配矩阵」（集成评审 Accept，overall 7.04，primary_index=2）；"
    "H-01（不确定性量化划可信度边界）、H-02（反事实解释干预）保留为 draft，未入选。"
    "本轮未因反对证据把某条标成 low 后自动重跑。"
    "入选后的真实问题在评审弱点，不在证据等级：验证写成万例前瞻交互日志与医疗纠纷发生率，数据难、统计效力不足。"
    "因此第二轮由人工从文献挖掘起 fork 重跑，要求把验证收到已有公开医学影像，并加强伦理/责任边界。"
)

P13_DOMAIN = (
    "案例所属的科学领域及问题特点：信息科学 / 医疗人工智能；问题特点为开放性问题（“人工智能能否取代医生？”），必须把口号收成可检验的辅助诊断与人机协同命题，而不是可立即完全取代医生的结论。"
)
P13_WHY = (
    "团队选择它进行完整展示的实际理由：该题完整跑通问题理解→文献挖掘→知识缺口→假设生成→假设评审；第一轮已冻结（run_id=c6e7fbfb-645a-461e-8e57-5ceb515bd86a）；"
    "第二轮用真实百炼环境从文献挖掘起 fork 重跑（run_id=f2c80a5f-0a67-44eb-b6bc-5770ed7d7a32），可对照两轮真实输出。"
)
P13_CAN = (
    "该案例能够展示本作品的哪些关键能力：事实与推断的区分（Fact 白名单）、缺口驱动假设生成、集成评审暴露可验证性弱点、HITL 暂停、人工反馈后的假设收敛。"
)
P13_CANT = (
    "该案例表现不能代表哪些题目或条件：不能代表真实前瞻临床试验或万例交互日志验证；不能代表纯数学/纯理论证明题；"
    "不能代表“证据弱自动迭代”路径——本案例三条假设均为 high，第二轮是人工 fork，不是低证据自动门禁。"
)

P14_EVIDENCE = (
    "第一轮使用了哪些证据和约束：题目背景段落 + 检索进入事实白名单的文献事实（facts 15、evidence 5、来源论文 4、入库文献 9）"
    " + 约束（《执业医师法》、须有不确定性/可解释输出、场景限定辅助诊断）。"
)
P14_OUTPUT = (
    "第一轮要求系统生成什么内容：候选假设（每条含形成依据、可检验预测、不确定性）+ 假设评审意见。"
    "评审后 HITL 暂停，Pipeline 的 ITERATIVE_EXPERIMENT 仍为 PENDING；实际实验在实验页完成。"
)
P14_MODEL = (
    "第一轮采用了哪些主要模型或方法设置：阶段记录为 qwen3.7-max（阿里云百炼）；七阶段流水线默认配置；"
    "知识库对齐门禁 + 集成评审；pause_after_hypothesis_review=true。"
)
P14_EVAL = (
    "第一轮结果使用什么口径评价：集成评审 Accept/分数（overall 7.04）+ 各假设新颖性/可检验性/数据可得性评分。"
    "三条假设 evidence_level 均为 high，本轮不以“低证据”作为失败口径。"
)

P16_TRIGGER = (
    "哪条证据、评价或人工意见触发了调整：集成评审弱点（需补充可验证细节；万例前瞻交互日志数据难、成本极高；"
    "医疗纠纷发生率低导致统计效力不足）+ 人工反馈。"
    "不是低证据自动门禁：三条假设均为 high，未点击“证据链迭代”按钮；"
    "第二轮为 POST /human-loop/rerun-from-stage，from_stage=literature_mining。"
)
P16_CHANGE = (
    "第二轮增加、删除或改变了什么：增加——打开 enable_gap_search，缺口补搜事实 9→10；"
    "改变——入选假设从 H-03（三维责任矩阵+前瞻）改为新 H-01（MC-Dropout 不确定性量化 + 公开影像可执行验证）；"
    "未删除第一轮实验记录，也未重跑迭代实验。"
)
P16_KEEP = (
    "哪些内容保持不变，为什么：评价口径（判定式门禁与集成评审阈值）不变，保证两轮可对比；"
    "核心方向不变——辅助诊断的可信度边界与人机协同责任，而不是完全取代医生。"
)
P16_EXPECT = (
    "团队预期第二轮在哪些方面改善：验证场景从不可执行的万例前瞻收到已有公开医学影像；"
    "可检验预测更具体；反对证据与伦理/责任边界更清楚。"
)

P17_IMPROVE = (
    "第二轮实际改善了什么：入选假设的验证收到 CheXpert/ISIC 等公开影像上的不确定性门控（ECE/F1 等可量化指标）；"
    "集成评审 Accept，overall 7.04→7.29；缺口补搜补充了动态评估指标/合成数据校准相关事实；"
    "表述上更明确：现行法律完善前模块只作过渡参考，最终责任仍由执业医师承担。"
)
P17_COST = (
    "哪些方面没有改善或出现了新的代价：真实前瞻临床仍未开展；第二轮没有重跑迭代实验，"
    "第一轮 RSNA 肺炎子集烟雾测试（250 行，silhouette 0.3219 等）不能当作新 H-01 的验证结果；"
    "Token 与耗时增加（文献约 77s、缺口约 225s、假设约 137s、评审约 347s）；"
    "仍有一条依赖约 200 名医师模拟交互的候选（新 H-03），只是未入选。"
)
P17_STOP = (
    "团队为什么在此时停止，或为什么仍需继续迭代：假设评审后再次 HITL 暂停，便于对照两轮假设，且不覆盖已成功的第一轮实验记录。"
    "若继续，应在实验页按新 H-01 设计公开影像验证，而不是把旧烟雾测试指标当作新假设的验证。"
)


def patch_template(doc: Document) -> None:
    paras = doc.paragraphs
    set_para(paras[73], P8)
    set_para(paras[85], P10)
    set_para(paras[107], P13_DOMAIN)
    set_para(paras[108], P13_WHY)
    set_para(paras[109], P13_CAN)
    set_para(paras[110], P13_CANT)
    set_para(paras[115], P14_EVIDENCE)
    set_para(paras[116], P14_OUTPUT)
    set_para(paras[117], P14_MODEL)
    set_para(paras[118], P14_EVAL)
    set_para(paras[128], P16_TRIGGER)
    set_para(paras[129], P16_CHANGE)
    set_para(paras[130], P16_KEEP)
    set_para(paras[131], P16_EXPECT)
    set_para(paras[136], P17_IMPROVE)
    set_para(paras[137], P17_COST)
    set_para(paras[138], P17_STOP)

    t13 = doc.tables[13]
    set_cell(t13.rows[1].cells[1], "人工智能能否取代医生？（sjtu_q_087，信息科学）")
    set_cell(
        t13.rows[2].cells[1],
        "研究对象：医疗AI辅助诊断的不确定性量化与人机协同责任分配（内：模型不确定性/可解释性；外：诊疗规范与侵权责任；边界：影像/病理辅助诊断，非自主治疗）。"
        "关键变量：预测不确定性、病例复杂度、医生经验、责任分配权重。",
    )
    set_cell(
        t13.rows[3].cells[1],
        "AI在部分任务上可超过人类，但实施障碍将长期阻碍大规模替代医生；部署须处理伦理与责任。"
        "第一轮事实白名单 15 条、证据原文 5、来源论文 4、入库文献 9（均绑定 source_chunk_id）。",
    )
    set_cell(
        t13.rows[4].cells[1],
        "gap_001：合成数据补稀缺高危场景时分布偏移会改变可信度边界，缺少人机协同责任评估标准；"
        "gap_002：高置信错误建议下，医生难以判断可信度边界，动态责任分配难以落地。",
    )
    set_cell(
        t13.rows[5].cells[1],
        "法律责任框架仍不明确；医学影像AI研究存在高偏倚风险与报告标准不足；"
        "不得把辅助诊断写成可立即完全取代医生。",
    )

    t14 = doc.tables[14]
    set_cell(
        t14.rows[1].cells[0],
        "H-01：在数据稀缺的高危临床场景中，基于不确定性量化的AI辅助诊断可界定可信度边界并支持人机协同，而非完全取代医生；落地须克服前瞻验证缺乏、高偏倚风险与伦理挑战。",
    )
    set_cell(
        t14.rows[1].cells[1],
        "fact_004, evfact_009, evfact_002, evfact_005, fact_001, evfact_004, fact_003, fact_005, paper_fact_005, paper_fact_006, fact_002",
    )
    set_cell(
        t14.rows[1].cells[2],
        "合成数据难拟合真实长尾；医生对不确定性分数的理解与接受度存在个体差异。",
    )
    set_cell(
        t14.rows[1].cells[3],
        "在含罕见/合成病例的影像与病历集上，用贝叶斯网络或 Deep Ensembles 输出不确定性，比较不同阈值下人机准确率与责任模拟。",
    )
    set_cell(t14.rows[1].cells[4], "draft；evidence_level=high；未入选")

    set_cell(
        t14.rows[2].cells[0],
        "H-02：在应对伦理挑战并克服前瞻验证缺乏的前提下，引入反事实解释的置信度干预，有望缓解自动化偏见、识别决策边界，促进人机协同而非完全取代医生。",
    )
    set_cell(
        t14.rows[2].cells[1],
        "evfact_009, fact_004, evfact_002, fact_001, evfact_004, fact_005, fact_003, paper_fact_005, paper_fact_006, fact_002",
    )
    set_cell(
        t14.rows[2].cells[2],
        "反事实解释在多模态高维空间可能缺乏临床合理性；实验样本量可能不足以覆盖复杂场景。",
    )
    set_cell(
        t14.rows[2].cells[3],
        "无解释 / 显著性图 / 反事实三组对照，记录准确率、决策时间、对AI建议的采纳率与责任归属判定。",
    )
    set_cell(t14.rows[2].cells[4], "draft；evidence_level=high；未入选")

    set_cell(
        t14.rows[3].cells[0],
        "H-03：构建基于AI预测不确定性、病例复杂度和医生经验水平的三维动态责任分配矩阵，为人机协同诊断提供责任划分框架；真实前瞻效果仍待验证。",
    )
    set_cell(
        t14.rows[3].cells[1],
        "fact_004, evfact_009, evfact_002, fact_001, evfact_004, fact_005, fact_003, paper_fact_005, paper_fact_006, fact_002",
    )
    set_cell(
        t14.rows[3].cells[2],
        "前瞻周期长、成本高；医疗纠纷发生率低可能导致统计效力不足；医生可能抵触系统自动分配责任。",
    )
    set_cell(
        t14.rows[3].cells[3],
        "合作医院前瞻观察：部署三维责任矩阵辅助诊断，收集诊断结果、责任日志、不良事件与纠纷发生率，倾向评分匹配对照。",
    )
    set_cell(
        t14.rows[3].cells[4],
        "accepted；evidence_level=high；集成评审 Accept，overall 7.04，primary_index=2",
    )

    t15 = doc.tables[15]
    set_cell(
        t15.rows[1].cells[1],
        "评审后 HITL 暂停，Pipeline 迭代实验阶段仍为 PENDING。"
        "实验页按入选 H-03 完成 1 组 1 轮沙箱（run_mode=smoke_only）：RSNA 肺炎影像子集，250 行×19 列；"
        "特征为 ai_entropy、complexity_score、reader_years_mean。",
    )
    set_cell(
        t15.rows[1].cells[2],
        "检验三维责任区间能否在公开影像烟雾测试中分开：轮廓系数 0.3219；"
        "不能替代万例前瞻或纠纷发生率验证。",
    )
    set_cell(
        t15.rows[2].cells[1],
        "KMeans 划分 AI 主导 / 协同 / 医生主导三区，比较各区 AUC 与 reader_agreement。",
    )
    set_cell(
        t15.rows[2].cells[2],
        "区分简单场景 AI 是否可靠、复杂高不确定场景医生是否更优："
        "AI 主导 AUC 0.8111，协同 0.9323，医生主导 0.9966；三区样本 69 / 112 / 69。",
    )
    set_cell(
        t15.rows[3].cells[1],
        "记录实验自身局限：总样本仅 250，统计效力受限；与评审弱点（万例前瞻难、纠纷发生率低）一致。",
    )
    set_cell(
        t15.rows[3].cells[2],
        "说明第一轮“计划层前瞻”与“落地层烟雾测试”的差距，作为第二轮收窄验证场景的依据。",
    )

    t17 = doc.tables[17]
    set_cell(
        t17.rows[1].cells[0],
        "入选 H-03 的验证依赖万例前瞻交互日志，评审写明数据难、成本极高",
    )
    set_cell(t17.rows[1].cells[1], "集成评审 weaknesses；H-03 可检验性原文")
    set_cell(t17.rows[1].cells[2], "研究计划在真实临床不可执行")
    set_cell(
        t17.rows[1].cells[3],
        "人工反馈要求收到已有公开医学影像，补充可落地的对照/校准指标",
    )
    set_cell(
        t17.rows[2].cells[0],
        "以医疗纠纷发生率为主终点，统计效力不足、验证周期漫长",
    )
    set_cell(t17.rows[2].cells[1], "集成评审 weaknesses")
    set_cell(t17.rows[2].cells[2], "计划难以证伪")
    set_cell(t17.rows[2].cells[3], "第二轮不再以纠纷发生率作为入选假设的主验证终点")
    set_cell(
        t17.rows[3].cells[0],
        "可验证细节不足；三条假设均围绕不确定性/责任/伦理，区分度有限。"
        "三条均为 high，不能写成系统因证据弱自动迭代",
    )
    set_cell(t17.rows[3].cells[1], "评审弱点 + evidence_level=high")
    set_cell(t17.rows[3].cells[2], "若按“低证据自动重跑”叙述则与事实不符")
    set_cell(
        t17.rows[3].cells[3],
        "未点证据链迭代按钮；人工 rerun-from-stage（literature_mining），打开 gap 补搜，要求加强反对证据与伦理约束",
    )

    t19 = doc.tables[19]
    set_cell(t19.rows[1].cells[1], "facts 15；未开 gap 补搜；enable_gap_search=False")
    set_cell(t19.rows[1].cells[2], "缺口补搜 facts 9→10（gap_literature_enrichment）")
    set_cell(t19.rows[1].cells[3], "人工打开 enable_gap_search=true，不是低证据自动门禁")
    set_cell(t19.rows[1].cells[4], "补充了动态评估指标 / 合成数据校准相关事实")
    set_cell(
        t19.rows[2].cells[1],
        "H-01/H-02 draft，H-03 accepted（三维责任矩阵）",
    )
    set_cell(
        t19.rows[2].cells[2],
        "新 H-01 accepted（MC-Dropout 不确定性+动态阈值，公开影像）；"
        "新 H-02 draft（MMD 分布偏移校准）；新 H-03 draft（交互式注意力，约 200 名医师模拟，未入选）",
    )
    set_cell(t19.rows[2].cells[3], "人工要求可执行验证，收窄到辅助诊断公开影像")
    set_cell(t19.rows[2].cells[4], "入选方向从“前瞻责任矩阵”收到“公开影像上的不确定性门控”")
    set_cell(t19.rows[3].cells[1], "集成评审 Accept，overall 7.04，primary=H-03")
    set_cell(t19.rows[3].cells[2], "集成评审 Accept，overall 7.29，入选新 H-01（评审分 7.8）")
    set_cell(t19.rows[3].cells[3], "同一套评审口径，输入因人工反馈与补搜而变化")
    set_cell(t19.rows[3].cells[4], "分数略升；入选假设更换")
    set_cell(
        t19.rows[4].cells[1],
        "实验页 1 组 1 轮 RSNA 肺炎子集烟雾测试成功（250 行；silhouette 0.3219 等）；Pipeline 实验阶段 PENDING",
    )
    set_cell(
        t19.rows[4].cells[2],
        "第二轮未重跑迭代实验；评审后再次 HITL 暂停（HUMAN_REVIEW_REQUIRED）",
    )
    set_cell(t19.rows[4].cells[3], "避免覆盖第一轮已成功实验记录；先对照假设")
    set_cell(
        t19.rows[4].cells[4],
        "实验指标仍属第一轮 H-03 烟雾测试，不能当作新 H-01 的验证结果",
    )
    set_cell(t19.rows[5].cells[1], "已限定辅助诊断，不得完全取代医生")
    set_cell(
        t19.rows[5].cells[2],
        "更明确：法律框架完善前仅作过渡参考，最终诊断与法律责任仍由执业医师承担",
    )
    set_cell(t19.rows[5].cells[3], "人工反馈第 3 条（收窄场景、禁止写成可立即临床部署）")
    set_cell(t19.rows[5].cells[4], "核心方向不变，表述更收紧")


def patch_huizong(doc: Document) -> None:
    p = doc.paragraphs
    set_para(p[97], "一项真实中间结果示例：" + P8.replace("\n", " "))
    set_para(p[112], "筛选结果如何形成：" + P10)
    set_para(p[132], "案例选择·" + P13_DOMAIN)
    set_para(p[133], "案例选择·" + P13_WHY)
    set_para(p[134], "案例选择·" + P13_CAN)
    set_para(p[135], "案例选择·" + P13_CANT)
    set_para(
        p[137],
        "案例内容1·科学问题原文：人工智能能否取代医生？（sjtu_q_087，信息科学）",
    )
    set_para(
        p[138],
        "案例内容2·研究对象与关键变量：内——医疗AI不确定性量化与可解释性；外——诊疗规范与侵权责任；"
        "边界——影像/病理辅助诊断。关键变量：预测不确定性、病例复杂度、医生经验、责任分配权重。",
    )
    set_para(
        p[139],
        "案例内容3·已有认识与主要证据：AI在部分任务上可超过人类，但实施障碍将长期阻碍大规模替代医生；"
        "部署须处理伦理与责任。第一轮 facts 15、evidence 5、来源论文 4、入库文献 9。",
    )
    set_para(
        p[140],
        "案例内容4·尚未解决的知识缺口：gap_001 合成数据偏移与责任评估缺失；"
        "gap_002 高置信错误下可信度边界与动态责任难落地。",
    )
    set_para(
        p[141],
        "案例内容5·需要保留的不确定性或争议：法律责任框架仍不明确；影像AI高偏倚风险与报告标准不足；"
        "不得把辅助诊断写成可立即完全取代医生。",
    )
    set_para(p[142], "第一轮设置·" + P14_EVIDENCE)
    set_para(p[143], "第一轮设置·" + P14_OUTPUT)
    set_para(p[144], "第一轮设置·" + P14_MODEL)
    set_para(p[145], "第一轮设置·" + P14_EVAL)
    set_para(
        p[147],
        "候选假设 H-01：数据稀缺高危场景下用不确定性量化划可信度边界，人机协同而非取代医生；"
        "依据 fact_001–005 等 11 个 fact_id；风险：合成数据难拟合长尾、医生接受度不一；"
        "可检验：BNNs/Deep Ensembles + 不同不确定性阈值；处理：draft / high，未入选。",
    )
    set_para(
        p[148],
        "候选假设 H-02：反事实解释干预以缓解自动化偏见、识别决策边界；"
        "风险：多模态反事实可能不合理、样本量可能不够；"
        "可检验：无解释/显著性图/反事实三组对照；处理：draft / high，未入选。",
    )
    set_para(
        p[149],
        "候选假设 H-03：不确定性×病例复杂度×医生经验的三维动态责任矩阵；"
        "风险：前瞻周期长、纠纷发生率低导致统计效力不足；"
        "可检验：合作医院前瞻观察；处理：accepted / high；集成评审 Accept，overall 7.04。",
    )
    set_para(
        p[150],
        "研究计划步骤1：评审后 HITL 暂停，Pipeline 迭代实验阶段 PENDING。"
        "实验页按 H-03 做 RSNA 肺炎子集烟雾测试（250×19）。"
        "支持/反对/区分：检验三维责任区间能否在公开影像上分开，不能替代万例前瞻。",
    )
    set_para(
        p[151],
        "研究计划步骤2：KMeans 划分 AI 主导/协同/医生主导，比较 AUC。"
        "结果：silhouette 0.3219；AUC 0.8111 / 0.9323 / 0.9966；样本 69 / 112 / 69。",
    )
    set_para(
        p[152],
        "研究计划步骤3：记录局限——总样本仅 250，统计效力受限；与评审弱点一致，作为第二轮收窄验证场景的依据。",
    )
    set_para(
        p[153],
        "说明：以上为 sjtu_q_087 第一轮冻结件真实结果（run_id=c6e7fbfb-…），不是 094 拟写稿。",
    )
    set_para(
        p[155],
        "问题1：入选 H-03 验证依赖万例前瞻交互日志；依据：集成评审 weaknesses；"
        "影响：计划不可执行；调整：人工要求收到已有公开医学影像。",
    )
    set_para(
        p[156],
        "问题2：以医疗纠纷发生率为主终点，统计效力不足；依据：评审弱点；"
        "影响：难以证伪；调整：第二轮不再以纠纷发生率作为入选假设主终点。",
    )
    set_para(
        p[157],
        "问题3：可验证细节不足、三条假设同质；且均为 high，不能写成证据弱自动迭代。"
        "调整：未点证据链迭代；人工 rerun-from-stage（literature_mining），打开 gap 补搜。",
    )
    set_para(p[158], "调整逻辑·" + P16_TRIGGER)
    set_para(p[159], "调整逻辑·" + P16_CHANGE)
    set_para(p[160], "调整逻辑·" + P16_KEEP)
    set_para(p[161], "调整逻辑·" + P16_EXPECT)
    set_para(
        p[163],
        "变化1·科学证据：第一轮 facts 15、未开 gap 补搜；第二轮缺口补搜 9→10；"
        "原因：人工打开 enable_gap_search；结果：补充动态评估/合成数据校准事实。",
    )
    set_para(
        p[164],
        "变化2·候选假设：第一轮 H-03 入选；第二轮新 H-01（MC-Dropout+公开影像）入选，"
        "新 H-02/H-03 为 draft；原因：人工要求可执行验证；结果：入选方向收窄。",
    )
    set_para(
        p[165],
        "变化3·假设比较与筛选：第一轮 Accept 7.04 / H-03；第二轮 Accept 7.29 / 新 H-01；"
        "原因：同一口径、输入变化；结果：分数略升，入选假设更换。",
    )
    set_para(
        p[166],
        "变化4·研究计划/实验：第一轮实验页 1 轮 RSNA smoke 成功；"
        "第二轮未重跑迭代实验，评审后再次 HITL 暂停。"
        "结果：实验指标仍属第一轮，不能当作新 H-01 的验证。",
    )
    set_para(
        p[167],
        "变化5·不确定性与边界：核心方向不变（辅助诊断、人机协同）；"
        "第二轮更明确最终责任仍由执业医师承担。",
    )
    set_para(p[168], "第二轮结论·" + P17_IMPROVE)
    set_para(p[169], "第二轮结论·" + P17_COST)
    set_para(p[170], "第二轮结论·" + P17_STOP)
    set_para(
        p[171],
        "说明：P13–P17 已按 sjtu_q_087 第一轮冻结件与 round2/ 真实运行改写；"
        "第二轮未重跑实验。P1 模型说明与 P20 交付入口按用户要求暂不改。",
    )


def main() -> None:
    tpl = Document(str(TPL))
    patch_template(tpl)
    tpl.save(str(TPL))
    print("patched template", TPL)

    hz = Document(str(HZ))
    patch_huizong(hz)
    hz.save(str(HZ))
    print("patched huizong", HZ)


if __name__ == "__main__":
    main()
