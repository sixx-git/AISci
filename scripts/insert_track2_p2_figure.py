# -*- coding: utf-8 -*-
"""把赛道二 P2 总体思路图插入正式模板，并更新待办标记。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from lxml import etree

ROOT = Path(r"d:/Workplace/AISci")
TPL = ROOT / "output/提交/模板/赛道二-方向1B-科研影响力分析与偏差解释-提交要求及模板.docx"
P2_PNG = ROOT / "output/innovation_schematics/fig_aisci_track2_p2_overview.png"
P6_PNG = ROOT / "output/innovation_schematics/fig_aisci_track2_p6_architecture.png"
TODO = ROOT / "output/提交/赛道二模板_待完成清单.docx"
HZ = ROOT / "output/提交/赛道二模板填写稿_汇总.docx"

P2_CAPTION = (
    "图 P2　科研影响力分析总体思路。实线为分析对象与目的→数据和资料获取→内容解析"
    "→影响力判断→主要因素解释→偏差检查→质量核验或人工反馈→结果与边界；"
    "虚线为核验失败时回到资料获取或内容解析，或转人工复核。"
    "缺证据时标注不确定，不编造。"
)
P6_CAPTION = (
    "图 P6　本作品实际架构。实线为分析对象→原始资料→提取信息→影响力结果"
    "→因素解释→偏差提示→质量反馈；虚线标明 Qwen、外部工具和人工注入位置，"
    "以及引用缺失、解析失败或无法回溯时的降级与复核回流。"
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _wipe_runs(p_elm) -> None:
    for child in list(p_elm):
        if child.tag == qn("w:r"):
            p_elm.remove(child)


def _underlined_run(text: str):
    run_elm = etree.Element("{%s}r" % W_NS)
    rpr = etree.SubElement(run_elm, "{%s}rPr" % W_NS)
    etree.SubElement(rpr, "{%s}color" % W_NS).set("{%s}val" % W_NS, "222222")
    etree.SubElement(rpr, "{%s}u" % W_NS).set("{%s}val" % W_NS, "single")
    etree.SubElement(rpr, "{%s}lang" % W_NS).set("{%s}eastAsia" % W_NS, "zh-CN")
    t = etree.SubElement(run_elm, "{%s}t" % W_NS)
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return run_elm


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = etree.Element("{%s}p" % W_NS)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para._p.append(_underlined_run(text))
    return new_para


def find_placeholder(doc: Document, needle: str, caption: str):
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    prefix = "图 P2" if caption.startswith("图 P2") else "图 P6"
    for p in doc.paragraphs:
        nxt = p._p.getnext()
        if nxt is None or nxt.tag != qn("w:p"):
            continue
        nxt_text = "".join(nxt.itertext())
        if nxt_text.startswith(prefix) and "w:drawing" in p._p.xml:
            return p
    raise RuntimeError(f"placeholder not found: {needle}")


def insert_figure_at_placeholder(doc: Document, needle: str, png: Path, caption: str) -> None:
    if not png.exists():
        raise FileNotFoundError(png)
    para = find_placeholder(doc, needle, caption)
    _wipe_runs(para._p)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(png), width=Cm(15.4))
    nxt = para._p.getnext()
    already = False
    if nxt is not None and nxt.tag == qn("w:p"):
        nxt_text = "".join(nxt.itertext())
        if nxt_text.startswith("图 P2" if caption.startswith("图 P2") else "图 P6"):
            already = True
            _wipe_runs(nxt)
            nxt.append(_underlined_run(caption))
    if not already:
        insert_paragraph_after(para, caption)


def patch_template() -> None:
    doc = Document(str(TPL))
    insert_figure_at_placeholder(doc, "请插入总体思路图", P2_PNG, P2_CAPTION)
    insert_figure_at_placeholder(doc, "请插入本作品实际架构图", P6_PNG, P6_CAPTION)
    doc.save(str(TPL))
    print("template saved", TPL.name, "inline", len(doc.inline_shapes))


def patch_notes() -> None:
    if TODO.exists():
        doc = Document(str(TODO))
        for p in doc.paragraphs:
            t = p.text or ""
            if t.startswith("[T2-01]"):
                _wipe_runs(p._p)
                p.add_run(
                    "[T2-01] P2 总体思路图：已插入正式模板 P2（fig_aisci_track2_p2_overview.png），"
                    "风格对齐赛道一 P6/P12。"
                )
            elif t.startswith("[T2-08]"):
                _wipe_runs(p._p)
                p.add_run(
                    "[T2-08] P6 架构图：已插入正式模板 P6（fig_aisci_track2_p6_architecture.png），"
                    "风格对齐赛道一 P6。"
                )
        doc.save(str(TODO))
        print("todo saved")
    if HZ.exists():
        doc = Document(str(HZ))
        for p in doc.paragraphs:
            t = p.text or ""
            if t.startswith("总体思路："):
                _wipe_runs(p._p)
                p.add_run(
                    "总体思路：已插入正式模板 P2。图文件 output/innovation_schematics/"
                    "fig_aisci_track2_p2_overview.png"
                )
            elif "P6 架构图" in t and "待处理" in t:
                _wipe_runs(p._p)
                p.add_run(
                    "P6 架构图：已插入正式模板 P6。图文件 output/innovation_schematics/"
                    "fig_aisci_track2_p6_architecture.png"
                )
        doc.save(str(HZ))
        print("huizong saved")


if __name__ == "__main__":
    patch_template()
    patch_notes()
