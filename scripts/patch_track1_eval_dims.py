# -*- coding: utf-8 -*-
"""按代码统一赛道一评价维度表述，并修正附录领域对齐公式。

代码依据：
- HypothesisReviewAgent：五维 scientific_value / novelty / testability /
  data_availability / cost_risk + ensemble_review（Accept 阈值 6.5）
- HypothesisTournamentSkill：两两比较四维——新颖性、可验证性、与证据一致性、可行性（辅助排序）
- hypothesis_generation_agent 偏题检测：有重叠 30+ratio*70；无重叠 max(5, n*2)；
  无关领域封顶 20；<30 判偏题
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from retune_track1_1a_theme import set_cell, set_fill  # noqa: E402

TPL = Path(
    r"d:/Workplace/AISci/output/提交/模板/赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx"
)

# P5 评价表：假设比较与选择（原误写为四项两两比较即正式入选）
TABLE3_R3 = [
    "假设比较与选择",
    "正式入选：Qwen 按科学价值、新颖性、可检验性、数据可得性、成本风险五维打分，"
    "再与多评审者集成（导师模拟、新颖性 Skill、证据规则）加权得综合分，≥6.5 为 Accept。"
    "锦标赛另按新颖性、可验证性、与证据一致性、可行性做辅助两两排序，不单独决定入选。",
    "同一轮全部候选。",
    "Accept 进入研究计划；未达阈值保留为草稿候选并写明未入选理由。",
]

# P10 正文：把锦标赛四维写清为辅助，与五维正式评审并列，避免只写五维却与 P1「锦标赛」脱节
P10 = (
    "筛选不是挑一条写得最完整的假设，而是留下可复核的比较轨迹：每条候选假设按固定顺序通过六道检查，"
    "任何一道不通过都在审计链中写明原因。"
    "①领域相关性——与研究问题领域无关的假设直接淘汰；"
    "②知识库对齐——假设必须绑定事实编号（fact_id），无法回溯事实的候选不进入后续；"
    "③引用校验——虚构或无法核验的引用被过滤；"
    "④可证伪性——写不出正反验证判据的候选降级；"
    "⑤去重——同质假设合并或淘汰；"
    "⑥正式评审——LLM 结构化五维打分（科学价值、新颖性、可检验性、数据可得性、成本风险），"
    "辅以锦标赛两两比较（新颖性、可验证性、与证据一致性、可行性）排序，"
    "再由多评审者集成（导师模拟、新颖性 Skill、证据规则加权）给出综合分；"
    "≥6.5 入选，低于阈值的保留为草稿候选并写明未入选理由。"
    "以 sjtu_q_087 为例：第一轮三条候选（H-01 不确定性量化、H-02 反事实解释、H-03 三维责任矩阵）"
    "证据等级均为 high，因此不是因证据不足被淘汰，而是按集成评审排序——"
    "H-03 综合分最高入选（Accept，overall 7.04），H-01/H-02 保留为草稿候选并写明未入选。"
    "入选后的真实问题出在计划层而非假设层：H-03 的验证写成上万例前瞻性临床研究与纠纷发生率统计，"
    "数据获取难、统计效力不足。系统如实记下该计划不可执行（创新二），"
    "没有把第一轮改写成一上来就选公开影像方案；第二轮才针对该弱点人工分支重跑，收窄验证场景。"
)

# 附录：补全对齐分分支，解释为何「低于 30 判偏题」与「有重叠时下限 30」不矛盾
APPENDIX_QUALITY = (
    "质量口径另计（在工程口径之外单独统计）："
    "① 引用核验——可疑引用 0 条，经 Fact 白名单与反证检索过滤虚构引用，所有引用均可回溯；"
    "② 假设-证据对齐率约 56.2%——假设能回溯到高证据等级事实（fact_id）的比例，"
    "其余约 43.8% 的假设未回溯到高证据等级事实，不能据此认为每条假设都已受充分证据约束；"
    "③ 领域对齐率约 79.1%——假设生成后由代码规则计算（非大模型自评）："
    "用人工内置的 6 个领域关键词集合，从研究问题与假设文本抽取主题词并求交集。"
    "计分：若研究问题抽不出主题词，对齐分记 50；"
    "若有关键词重叠，对齐分 = int(30 + 重叠率 × 70)（该分支下分值落在 30–100）；"
    "若无重叠，对齐分 = max(5, 假设主题词数 × 2)（可低于 30）；"
    "若命中明显无关领域关键词，对齐分再封顶为 20。"
    "对齐分 < 30 判为偏题并过滤。"
    "领域对齐率 = 未偏题假设数 / 全部候选假设数；"
    "④ 报告逻辑审查均分约 6.08/10——对输出报告逻辑自洽性的审查均分（满分 10 分）。"
)


def main() -> None:
    doc = Document(str(TPL))

    t3 = doc.tables[3]
    for ci, val in enumerate(TABLE3_R3):
        set_cell(t3.rows[3].cells[ci], val)

    # P10 正式填写段：当前在「请说明…」占位之后
    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        if t.startswith("筛选不是挑一条写得最完整的假设"):
            set_fill(p, P10)
            print("patched P10 para", i)
            break
    else:
        raise RuntimeError("P10 fill paragraph not found")

    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        if t.startswith("质量口径另计"):
            set_fill(p, APPENDIX_QUALITY)
            print("patched appendix para", i)
            break
    else:
        raise RuntimeError("appendix quality paragraph not found")

    # P19 结果表若仍写「四项两两」一并核对（当前应无）
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                tx = cell.text or ""
                if "两两比较" in tx and "四项" in tx and "正式" not in tx:
                    print("WARN leftover", f"T{ti}R{ri}C{ci}", tx[:80])

    out = TPL
    try:
        doc.save(str(out))
    except PermissionError:
        out = TPL.with_name(TPL.stem + "._eval_patched.docx")
        doc.save(str(out))
        print("LOCKED, wrote", out.name, "- close Word then re-run or replace manually")
    else:
        print("saved", out.name)

    # verify
    doc2 = Document(str(out))
    print("T3R3C1:", doc2.tables[3].rows[3].cells[1].text[:120])
    for p in doc2.paragraphs:
        if "对齐分 = int(30" in (p.text or "") or "无重叠，对齐分" in (p.text or ""):
            print("appendix ok:", (p.text or "")[200:360])
            break
    for p in doc2.paragraphs:
        if "辅以锦标赛两两比较" in (p.text or ""):
            print("P10 ok")
            break


if __name__ == "__main__":
    main()
