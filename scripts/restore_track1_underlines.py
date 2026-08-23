# -*- coding: utf-8 -*-
"""Restore template fill-in underline: label (no u) + answer (single underline, color 222222)."""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

TPL = Path(r"d:\Workplace\AISci\output\提交\模板\赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx")
HZ = Path(r"d:\Workplace\AISci\output\提交\赛道一模板填写稿_汇总.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _rpr(color: str, underline: bool) -> etree._Element:
    rpr = etree.Element("{%s}rPr" % W_NS)
    etree.SubElement(rpr, "{%s}color" % W_NS).set("{%s}val" % W_NS, color)
    if underline:
        etree.SubElement(rpr, "{%s}u" % W_NS).set("{%s}val" % W_NS, "single")
    etree.SubElement(rpr, "{%s}lang" % W_NS).set("{%s}eastAsia" % W_NS, "zh-CN")
    return rpr


def _make_run(text: str, underline: bool, color: str = "222222") -> etree._Element:
    r = etree.Element("{%s}r" % W_NS)
    r.append(_rpr(color, underline))
    t = etree.SubElement(r, "{%s}t" % W_NS)
    if text.startswith(" ") or text.endswith(" ") or "\n" in text:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def _wipe_runs(p_elm) -> None:
    for child in list(p_elm):
        if child.tag == qn("w:r"):
            p_elm.remove(child)


def _para_text(p) -> str:
    return "".join((r.text or "") for r in p.runs)


def restore_para(p, *, whole_fill_if_no_label: bool) -> str:
    """Return action tag or empty if skipped."""
    style = p.style.name if p.style else ""
    if style.startswith("Heading"):
        return ""
    raw = _para_text(p)
    if not raw.strip():
        return ""
    if raw.startswith("====") or raw.startswith("□"):
        return ""

    r0_u = bool(p.runs and p.runs[0].underline)
    colon = raw.find("：")
    p_elm = p._p

    # List / prompt+answer: first run historically has no underline.
    if colon >= 0 and (not r0_u or style.startswith("List")):
        label = raw[: colon + 1]
        answer = raw[colon + 1 :]
        # Don't split if the colon is not a prompt (e.g. only in the body and r0 already underlined)
        if r0_u and not style.startswith("List") and not _looks_like_prompt_label(label):
            label, answer = "", raw
        if answer.strip() == "" and label:
            return ""
        _wipe_runs(p_elm)
        if label:
            p_elm.append(_make_run(label, underline=False))
            if answer:
                p_elm.append(_make_run(answer, underline=True))
            return "split"
        _wipe_runs(p_elm)
        p_elm.append(_make_run(raw, underline=True))
        return "whole"

    # Whole-paragraph fill (instruction lives in previous heading)
    _wipe_runs(p_elm)
    p_elm.append(_make_run(raw, underline=True))
    return "whole"


PROMPT_PREFIXES = (
    "本作品针对的具体问题",
    "本作品实际完成的核心方法",
    "最有代表性的结果",
    "目前仍存在的主要局限",
    "本作品最终形成的主要输出",
    "本作品已使用官方",
    "相较直接使用大模型回答",
    "本作品如何判断来源是否可用",
    "遇到证据不足或相互冲突时",
    "本作品如何保留引用和科学证据之间的对应关系",
    "团队重点评价的是哪些方面",
    "评价由程序、模型、团队成员",
    "第一轮与第二轮是否采用同一口径",
    "什么情况下继续生成、停止或交由研究者判断",
    "团队为减少无依据生成",
    "上下文更新或多轮调用",
    "该设计对结果带来的实际变化",
    "案例所属的科学领域及问题特点",
    "团队选择它进行完整展示的实际理由",
    "该案例能够展示本作品的哪些关键能力",
    "该案例表现不能代表哪些题目或条件",
    "第一轮使用了哪些证据和约束",
    "第一轮要求系统生成什么内容",
    "第一轮采用了哪些主要模型或方法设置",
    "第一轮结果使用什么口径评价",
    "哪条证据、评价或人工意见触发了调整",
    "第二轮增加、删除或改变了什么",
    "哪些内容保持不变，为什么",
    "团队预期第二轮在哪些方面改善",
    "第二轮实际改善了什么",
    "哪些方面没有改善或出现了新的代价",
    "团队为什么在此时停止",
    "团队如何检查研究计划是否具体、可复核",
    "研究计划中哪些内容已经具备执行条件",
    "本作品如何避免用",
    "哪些反馈由系统自动产生",
    "哪些调整来自研究者或团队成员",
    "团队实际保留了哪些前后变化",
    "科学逻辑方面，本作品实际改善了什么",
    "技术方法方面，本作品实际改善了什么",
    "结果表现方面，本作品实际改善了什么",
    "没有改善的部分及增加的成本",
    "一项真实中间结果示例",
    "筛选结果如何形成",
    "案例选择·",
    "案例内容",
    "第一轮设置·",
    "候选假设 H-",
    "研究计划步骤",
    "问题1",
    "问题2",
    "问题3",
    "调整逻辑·",
    "变化1",
    "变化2",
    "变化3",
    "变化4",
    "变化5",
    "第二轮结论·",
    "说明：",
)


def _looks_like_prompt_label(label: str) -> bool:
    s = label.rstrip("：")
    return any(s.startswith(x) or x in s for x in PROMPT_PREFIXES) or (
        4 <= len(s) <= 40
    )


def patch_template(doc: Document) -> dict:
    stats = {"split": 0, "whole": 0, "skip": 0}
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        raw = _para_text(p)
        if not raw.strip() or style.startswith("Heading") or raw.startswith("[请") or raw.startswith("□"):
            stats["skip"] += 1
            continue
        if _already_ok(p):
            stats["skip"] += 1
            continue
        extra_empty = len(p.runs) > 1 and any((r.text or "") == "" for r in p.runs[1:])
        mashed_list = style.startswith("List") and "：" in raw
        mashed_normal = (
            (not style.startswith("List"))
            and extra_empty
            and "：" in (p.runs[0].text or "")
            and not p.runs[0].underline
        )
        whole_broken = extra_empty and bool(p.runs) and bool(p.runs[0].underline)
        if not (extra_empty or mashed_list or mashed_normal or whole_broken):
            stats["skip"] += 1
            continue
        act = restore_para(p, whole_fill_if_no_label=True)
        stats[act or "skip"] += 1
    return stats


def _already_ok(p) -> bool:
    runs = [r for r in p.runs if (r.text or "") != ""]
    if len(runs) == 2 and not runs[0].underline and runs[1].underline and "：" in (runs[0].text or "") and (runs[0].text or "").endswith("："):
        return True
    if len(runs) == 1 and runs[0].underline and "：" not in (runs[0].text or "")[:20]:
        return True
    if len(runs) >= 2 and not runs[0].underline and all(r.underline for r in runs[1:]) and (runs[0].text or "").endswith("："):
        # already correct, maybe fragmented answer runs like P90 — leave them
        if all((r.text or "") != "" for r in runs[1:]):
            return True
    return False


def patch_huizong(doc: Document) -> dict:
    stats = {"split": 0, "whole": 0, "skip": 0}
    for p in doc.paragraphs:
        raw = _para_text(p)
        if not raw.strip() or raw.startswith("===="):
            stats["skip"] += 1
            continue
        if "：" not in raw:
            stats["skip"] += 1
            continue
        act = restore_para(p, whole_fill_if_no_label=False)
        stats[act or "skip"] += 1
    return stats


def main() -> None:
    tpl = Document(str(TPL))
    s1 = patch_template(tpl)
    tpl.save(str(TPL))
    print("template", s1)

    hz = Document(str(HZ))
    s2 = patch_huizong(hz)
    hz.save(str(HZ))
    print("huizong", s2)


if __name__ == "__main__":
    main()
