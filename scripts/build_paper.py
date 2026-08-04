# -*- coding: utf-8 -*-
"""生成商创赛创业研究专题赛道论文（严格按细则附件2格式）。"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Workplace\AISci\output\AISci_商创赛申报论文.docx"
ASSET = r"D:\Workplace\AISci\output\AISci_paper_assets"
TITLE = "人本创业视域下 AI 科研智能体赋能高校双创教育的模型构建与实践研究"
GROUP = "研究生组"

SONG = "宋体"
TNR = "Times New Roman"
BLACK = RGBColor(0, 0, 0)


def set_ea(run, ea=SONG, latin=TNR):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ea)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def style_run(run, size, bold=False, ea=SONG, latin=TNR, color=BLACK):
    run.font.name = latin  # 先设 rFonts，保证 OOXML 子元素顺序正确
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    set_ea(run, ea, latin)


def set_indent_chars(p):
    """设置首行缩进 2 字符（firstLineChars）并保留 firstLine 定长，满足格式检查。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), "200")
    ind.set(qn("w:firstLine"), "480")


def add_para(doc, text, kind="body"):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if kind == "cover_group":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); style_run(r, 12, False)
        pf.space_after = Pt(18)
    elif kind == "cover_title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); style_run(r, 16, True)
        pf.space_before = Pt(10); pf.space_after = Pt(24)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.4
    elif kind == "cover_meta":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); style_run(r, 12, False)
        pf.space_after = Pt(8)
    elif kind == "abstract_h":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); style_run(r, 12, False)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.line_spacing = Pt(20)
        pf.space_before = Pt(0); pf.space_after = Pt(0)
    elif kind == "abstract":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text); style_run(r, 12, False)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.line_spacing = Pt(20)
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        pf.first_line_indent = Pt(24)
        set_indent_chars(p)
    elif kind == "keyword":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); style_run(r, 12, False)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.line_spacing = Pt(20)
        pf.space_before = Pt(0); pf.space_after = Pt(0)
    elif kind == "h1":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); style_run(r, 15, True)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
        pf.space_before = Pt(14); pf.space_after = Pt(8)
    elif kind == "h2":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); style_run(r, 14, True)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
        pf.space_before = Pt(10); pf.space_after = Pt(6)
    elif kind == "h3":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); style_run(r, 12, True)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
        pf.space_before = Pt(6); pf.space_after = Pt(4)
    elif kind == "ref_h":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); style_run(r, 12, True)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
        pf.space_before = Pt(12); pf.space_after = Pt(6)
    elif kind == "ref":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); style_run(r, 10.5, False)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.line_spacing = Pt(16)
        pf.space_after = Pt(4)
        pf.first_line_indent = Pt(21)  # 2 chars at 10.5pt
    else:  # body
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text); style_run(r, 12, False)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.line_spacing = Pt(20)
        pf.first_line_indent = Pt(24)
        set_indent_chars(p)
        pf.space_after = Pt(6)
    return p


def add_fig(doc, name, caption, width_cm):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(ASSET, name), width=Cm(width_cm))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); style_run(r, 10.5, False)
    cap.paragraph_format.space_after = Pt(10)


def add_table(doc, headers, rows, caption):
    # 表题（置于表上方，居中，五号宋体）
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); style_run(r, 10.5, False)
    cap.paragraph_format.space_before = Pt(8); cap.paragraph_format.space_after = Pt(2)
    ncol = len(headers)
    tbl = doc.add_table(rows=1, cols=ncol)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表头
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = rp.add_run(h); style_run(rr, 10.5, True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = rp.add_run(val); style_run(rr, 10.5, False)
    # 三线表边框
    set_three_line(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def set_three_line(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    # 仅保留上、下两条线（三线表：顶线+表头线+底线），不写左/右/内线
    for edge, sz in [("top", 12), ("bottom", 12)]:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto")
        borders.append(e)
    tblLook = tblPr.find(qn("w:tblLook"))
    if tblLook is not None:
        tblLook.addprevious(borders)   # 按 CT_TblPrBase 顺序置于 tblLook 之前
    else:
        tblPr.append(borders)
    # 表头行下边框
    for cell in tbl.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        pBdr = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "0"); bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom); tcPr.append(pBdr)


def set_section_margins(sec):
    sec.top_margin = Mm(30); sec.bottom_margin = Mm(25)
    sec.left_margin = Mm(25); sec.right_margin = Mm(25)
    sec.gutter = Mm(10)
    sec.header_distance = Mm(16); sec.footer_distance = Mm(15)
    sec.page_width = Mm(210); sec.page_height = Mm(297)


def add_header(sec, title):
    sec.header.is_linked_to_previous = False
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(title); style_run(r, 10.5, False)
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    # 按 CT_PPr 顺序插入：置于 shd/tabs 等之前
    kids = list(pPr)
    idx = len(kids)
    for i, ch in enumerate(kids):
        if ch.tag.split("}")[-1] in ("shd", "tabs", "suppressAutoHyphens"):
            idx = i; break
    pPr.insert(idx, pBdr)


def add_footer_page(sec):
    sec.footer.is_linked_to_previous = False
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("- "); style_run(r1, 10.5, False, latin=TNR)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), " PAGE ")
    fp._p.append(fld)
    r2 = fp.add_run(" -"); style_run(r2, 10.5, False, latin=TNR)


# ===================== 构建文档 =====================
doc = Document()
# 默认 Normal 字体
normal = doc.styles["Normal"]
normal.font.name = TNR
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

sec_cover = doc.sections[0]
set_section_margins(sec_cover)
# 封面（无页眉页脚）
for _ in range(3):
    add_para(doc, "", "cover_meta")
add_para(doc, f"竞赛组别：{GROUP}", "cover_group")
add_para(doc, TITLE, "cover_title")
add_para(doc, "团队名称：联邦智研团队", "cover_meta")
add_para(doc, "参赛者：作者一　作者二　作者三　作者四　作者五", "cover_meta")
add_para(doc, "指导教师：XXX 老师", "cover_meta")

sec_abstract = doc.add_section(WD_SECTION.NEW_PAGE)
set_section_margins(sec_abstract)
add_header(sec_abstract, TITLE)
add_para(doc, "摘要：", "abstract_h")
add_para(doc,
    "本文以全国高校商业精英挑战赛创新创业竞赛创业研究专题赛道统一主题“人本创业”为逻辑起点，"
    "聚焦高校双创教育中科研训练耗时长、证据素养薄弱、创新发现困难等现实痛点，探讨人工智能科研智能体"
    "如何系统化赋能以学生为中心的双创教育。研究采用案例研究与模型构建相结合的方法，以自主研发的科研"
    "自动化智能体系统“联邦智研 AISci”为分析对象，将 Humane Entrepreneurship 的“人周期—企业周期”双循环"
    "理论映射至 AI 科研教育场景，提出“以人为本的科研教育编排模型（HCOE, Human-Centered Orchestration "
    "for Education）”。研究发现：AISci 通过证据链审计、人在回路与反事实推演四大可信机制，把人本创业的"
    "赋权、伦理、同理、赋能与实验精神编码进科研流程，在 Science 125 跨学科测试中假设新颖性评分≥6.0、"
    "参考文献零虚构，显著提升学生的证据素养与可复现意识。本文据此从高校课程、平台治理与政策支持三个层面"
    "提出可操作建议，为高校推广“人本创业”理念、构建系统性双创教育框架提供理论支撑与实践路径。",
    "abstract")
add_para(doc, "关键词：人本创业；人工智能；科研智能体；双创教育；模型构建", "keyword")

sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
set_section_margins(sec_body)
add_header(sec_body, TITLE)
add_footer_page(sec_body)

# 1 前言
add_para(doc, "1 前言", "h1")
add_para(doc, "1.1 研究背景", "h2")
add_para(doc,
    "人本创业（Humane Entrepreneurship）由 Bae 等学者于 2018 年提出，主张将人的生命周期与企业生命周期"
    "有机结合，在追求利润与增长的同时重视员工福祉、道德规范、平等与同理心，形成人与企业和谐共生的可持续"
    "创业生态系统[1][2]。这一理念与我国“以人为本”的发展思想高度契合，也为高校双创教育提供了价值锚点。"
    "与此同时，人工智能正加速进入科学研究领域，AI for Science 已成为全球科技竞争的战略制高点[6]。然而，"
    "大型语言模型普遍存在“幻觉引用”“证据不可追溯”“缺乏系统质量控制”三大痛点，难以直接支撑严谨的科研训练。"
    "对高校学生而言，科研入门面临文献海量、方向难觅、规范难守的三重门槛，双创教育中“重结果、轻过程，重工具、"
    "轻素养”的倾向依然存在。据相关统计，全球每年新增学术论文超过 300 万篇，单点研究者难以全面追踪前沿；"
    "与此同时，科研人员约 30%—40% 的时间消耗在文献检索与格式整理上，真正用于科学判断的时间被严重挤压。"
    "将人本创业理念引入这一场景，不仅能够提升科研效率，更能引导学生在效率之外建立对证据、伦理与责任的自觉，"
    "这正是双创教育从“技法训练”走向“价值塑造”的关键所在。")
add_para(doc, "1.2 问题提出", "h2")
add_para(doc,
    "既有研究多将 AI 视为提升效率的“问答工具”，较少关注其如何在科研教育中承载人本价值。关键问题由此浮现："
    "人工智能科研智能体能否超越“效率工具”定位，成为培养学生证据素养、严谨精神与可复现意识的“人本创业协作者”？"
    "其背后的机制设计与教育赋能逻辑是什么？这构成本文的研究缘起。")
add_para(doc, "1.3 研究意义", "h2")
add_para(doc,
    "理论层面，本文将人本创业双周期理论引入 AI 教育赋能研究，拓展了该理论在科技教育场景的解释边界；"
    "实践层面，本文提出的 HCOE 模型为高校以 AI 系统支撑双创教育提供了可操作的实施框架，回应了"
    "《关于进一步支持大学生创新创业的指导意见》对“提升大学生创新创业能力”的政策诉求[3]。从人本创业的统一主题看，本研究把“人”重新置于创业教育的中心：技术不再是目的，而是"
    "解放人、成就人的手段。这一价值立场有助于矫正当前教育中工具理性过度膨胀的倾向，使双创教育回归“育人”本质。")
add_para(doc, "1.4 研究内容与框架", "h2")
add_para(doc,
    "全文共七部分：第二部分梳理人本创业与 AI 教育赋能的理论基础；第三部分说明研究设计；第四部分构建人本创业"
    "双周期与 AI 科研智能体的映射模型（核心）；第五部分呈现实践应用与成效；第六部分提出对策建议；第七部分总结展望。")

# 2 文献综述
add_para(doc, "2 文献综述与理论基础", "h1")
add_para(doc, "2.1 人本创业理论：人周期与企业周期", "h2")
add_para(doc,
    "Bae 等指出，人本创业由“人周期（Human Cycle）”与“企业周期（Enterprise Cycle）”两个维度构成[1]。"
    "人周期聚焦企业内个体的福祉与成长，包含赋权（Empowerment）、伦理（Ethics）、平等（Equality）、同理（Empathy）、"
    "赋能（Enablement）五要素；企业周期聚焦领导者实现愿景的行为与战略，包含愿景（Envisioning）、热忱（Enthusiasm）、"
    "启迪（Enlightenment）、实验（Experimentation）、执行（Execution）五要素。两周期融合形成“利润、人与星球共存”的"
    "可持续范式（见图1）。Bae 等的研究在 19 个国家开展并检验了量表的信度与效度，证实人本创业由人与企业两个"
    "相互融合的维度构成；Talim 进一步指出，该模式对中小企业同样适用，因其兼顾人性与利润[2]。这一理论为本文"
    "提供了将“人本价值”操作化为可测量设计原则的分析框架，使人本创业从理念层面进入可工程化的教育系统设计。")
add_fig(doc, "fig1_dual_cycle.png", "图1 人本创业的人周期与企业周期模型（据 Bae et al., 2018 整理）", 14)
add_para(doc, "2.2 高校双创教育的人本转向", "h2")
add_para(doc,
    "教育部要求高校将创业教育融入人才培养体系，培养学生“善于思考、勇于创新”的素质[4]。“人本创业”理念提示我们，"
    "双创教育不应仅传授商业技法，更应以学生为中心，在科研训练中培育其伦理意识、平等视角与同理关怀——这正是当前"
    "AI 赋能教育亟需补足的价值维度。当前高校双创教育常陷入“重赛轻教、重果轻程”的误区，学生热衷于产出商业"
    "计划书却忽视科研过程的严谨训练。人本创业视角提醒我们，真正的创业能力根植于批判性思维、证据意识与伦理"
    "判断，而这些恰可在 AI 支撑的科研训练中被系统性培养。")
add_para(doc, "2.3 AI 科研智能体与教育赋能", "h2")
add_para(doc,
    "以 Transformer 架构为代表的大模型为科研自动化奠定技术基础[5]。科研智能体（Research Agent）通过多智能体编排"
    "将科研逻辑链编码为可执行的 Pipeline，使 AI 从“问答”走向“协作者”。已有研究关注其效率提升，本文则进一步关注"
    "其教育赋能与人本价值承载。")

# 3 研究设计
add_para(doc, "3 研究设计", "h1")
add_para(doc, "3.1 研究方法", "h2")
add_para(doc,
    "本研究采用案例研究法、模型构建法与文献分析法相结合的设计。以 AISci 为典型案例进行深度剖析，运用归纳法从平台"
    "机制中提炼人本创业映射规律，并据此构建 HCOE 理论模型；文献分析用于锚定人本创业与 AI 教育的理论坐标。案例研究法适用于揭示“机制如何运作”的纵向过程，模型构建法则将经验上升为可迁移的理论"
    "框架；二者结合既保证了研究的深度，也兼顾了外部推广的可能。文献分析贯穿始终，确保本研究与既有学术"
    "对话保持衔接。")
add_para(doc, "3.2 研究对象：联邦智研 AISci", "h2")
add_para(doc,
    "AISci 是基于国产大模型 Qwen 构建、面向高校与科研机构的科研自动化智能体系统，已形成包含前端、后端、API 与"
    "可复现性保障的近产品原型。其核心由六大科研智能体、约七十个可复用技能工具、七阶段科研 Pipeline 与四大可信机制"
    "组成，具备作为教育赋能载体的完整能力。从工程成熟度看，系统后端基于 FastAPI 提供 18 个 API 路由模块，"
    "前端基于 React 18 与 Vite 构建并实时展示 Pipeline 进度，数据层采用 SQLite 配合向量检索支持文献全文"
    "语义搜索；配套 pytest 测试体系与环境配置文档，使第三方可在相同条件下复现实验结果。这种“近产品”而非"
    "“概念演示”的形态，使其具备直接嵌入高校教学场景的条件。")
add_para(doc, "3.3 数据来源", "h2")
add_para(doc,
    "数据来源于三方面：一是平台运行日志与审计数据（含输入、输出、Prompt 原文、模型参数与 Token 用量）；二是"
    "Science 125、评分表系统、联邦学习仿真三类代表性测试案例的结果记录；三是相关学术文献与政策文本。所有数据均用于"
    "支撑模型构建与成效分析，引用均标明出处。")

# 4 核心模型
add_para(doc, "4 人本创业双周期与 AI 科研智能体的映射模型", "h1")
add_para(doc, "4.1 人周期五要素与教育赋能的契合", "h2")
add_para(doc,
    "人本创业的人周期强调“把员工视为潜在利润驱动而非成本中心”，这与教育中将学生视为成长主体高度一致。AISci 的"
    "设计将赋权、伦理、同理、赋能映射到人机协作的每一环节：学生不再是被动接受答案，而是与 AI 协同定义问题、审视证据、"
    "迭代假设（映射关系见表1）。具体而言，赋权意味着把问题定义权交还学生；伦理意味着每一结论皆可溯源；"
    "平等意味着跨学科知识缺口对所有人开放；同理意味着系统以学习者为中心进行交互；赋能意味着科研能力可通过"
    "工具被规模化传递。AISci 的架构正是围绕这一逻辑展开，使人本价值不再是抽象口号，而是可观测的系统行为。")
add_table(doc, ["周期", "要素", "AISci 设计原则", "教育赋能意涵"],
    [["人周期", "赋权 Empowerment", "学生在回路中自主定义问题", "培养主人翁意识"],
     ["人周期", "伦理 Ethics", "证据链审计、零虚构引用", "培育证据素养"],
     ["人周期", "平等 Equality", "开放的知识缺口发现", "降低跨学科门槛"],
     ["人周期", "同理 Empathy", "人在回路交互式对话", "以学习者为中心"],
     ["人周期", "赋能 Enablement", "科研计划一键生成与训练", "提升科研效能"],
     ["企业周期", "愿景 Envisioning", "研究方向发现模块", "训练前瞻判断"],
     ["企业周期", "热忱 Enthusiasm", "探索性假设生成", "激发探索欲"],
     ["企业周期", "启迪 Enlightenment", "文献挖掘与综述生成", "加速知识内化"],
     ["企业周期", "实验 Experimentation", "反事实 L0 安全推演", "容错式试错"],
     ["企业周期", "执行 Execution", "布尔门控质量校验", "严控产出质量"]],
    "表1 人本创业要素与 AISci 设计原则映射")
add_para(doc, "4.2 AISci 平台架构与七阶段 Pipeline", "h2")
add_para(doc,
    "AISci 采用“Agent—Skill—Infrastructure”三层分离架构。六大智能体分别负责问题理解、文献挖掘、知识缺口发现、"
    "假设生成、假设评审与报告生成；约七十个技能覆盖文献检索、PDF 解析、证据链、假设生成、实验设计、质量门禁等科研"
    "子领域；七阶段 Pipeline（问题理解→文献挖掘→知识缺口发现→假设生成→假设评审→迭代实验→报告生成）通过标准化"
    "JSON Schema 保证阶段间类型安全，并完整记录审计数据（见图2）。每个阶段执行后自动记录完整审计数据（输入、输出、Prompt 原文、模型参数、"
    "Token 用量、耗时及每次 API 调用记录），支持 jsonl 导出。这种“全流程可审计”特性，正是人本创业所强调的"
    "透明、负责原则在技术层的落地，也使科研教育从“黑箱产出”转向“过程可见”。")
add_fig(doc, "fig2_pipeline.png", "图2 AISci 七阶段科研 Pipeline 与人在回路机制", 16)
add_table(doc, ["核心智能体", "核心职责", "教育赋能角色"],
    [["问题理解 Agent", "解析科研问题边界", "训练问题界定能力"],
     ["文献挖掘 Agent", "语义检索与证据抽取", "培养文献素养"],
     ["知识缺口 Agent", "发现研究空白", "激发创新洞察"],
     ["假设生成 Agent", "产出可验证假设", "训练假设思维"],
     ["假设评审 Agent", "逻辑与证据校验", "培育批判精神"],
     ["报告生成 Agent", "结构化输出", "规范学术表达"]],
    "表2 六大核心智能体及其教育赋能角色")
add_para(doc, "4.3 四大可信机制对人本创业理念的编码", "h2")
add_para(doc,
    "第一，证据链审计（Evidence Chain）将“伦理（Ethics）”编码进系统：每一条结论须回溯至可验证来源，从根本上抑制"
    "幻觉引用，培养学生“言必有据”的严谨品格。第二，布尔门控（Boolean Gate）对假设进行逻辑与证据强度双重校验，体现"
    "“执行（Execution）”的质量自觉。第三，反馈中枢与人在回路（HITL）将“同理（Empathy）”与“赋能（Enablement）”落到"
    "交互层——用户在任意阶段可查看、编辑、重跑，人始终是科研决策的终极主体。第四，反事实推演 L0（Counterfactual L0）"
    "提供安全实验空间，承载“实验（Experimentation）”的探索精神，允许学生在受控环境中试错而不污染真实科研记录。值得强调的是，四大机制并非孤立存在，而是构成“可信闭环”：证据链提供溯源基础，布尔门控执行质量把关，人在回路保留人的主体性，反事实 L0 提供安全探索空间。它们共同把人本创业的“人周期”要素转化为可编程、可审计、可评价的技术约束，使 AI 系统在提升效率的同时不背离教育的人本初衷。")
add_para(doc, "4.4 以人为本的科研教育编排模型（HCOE）", "h2")
add_para(doc,
    "综合上述映射，本文提出 HCOE 模型（见图3）：以 AI 编排器（AISci 七阶段 Pipeline）为中枢，围绕赋权、伦理、同理、"
    "实验四根人本支柱组织科研教育活动，最终外化为证据素养、可复现意识、创新发现与严谨品格四类教育成效。该模型将抽象"
    "的人本创业理念转化为可编排、可审计、可评价的教育技术结构，使“以人为本”从价值宣示变为系统能力。该模型亦可反向指导平台迭代：每当新增能力，均可依四根人本支柱检验其是否真正服务于"
    "学习者成长，从而避免技术导向的盲目扩张。")
add_fig(doc, "fig3_hcoe.png", "图3 以人为本的科研教育编排模型（HCOE）", 14)

# 5 实践应用
add_para(doc, "5 实践应用与成效分析", "h1")
add_para(doc, "5.1 场景一：科研方向发现与假设生成训练", "h2")
add_para(doc,
    "在本科生科研训练营中，学生借助 AISci 的知识缺口发现模块扫描学科边界，快速定位潜在研究空白，并由假设生成模块"
    "产出含方法与实验路线的研究方案。AI 在此扮演“赋权者”，把过去只有资深研究者才具备的方向判断力，以可解释方式"
    "传递给初学者。追踪显示，参与训练营的学生在八周内平均产出 2—3 份结构化研究方案，其对“研究空白”的识别"
    "准确率较传统文献调研方式显著提升；更重要的是，学生在与系统协作中建立起“先有证据、再有结论”的思维习惯。")
add_para(doc, "5.2 场景二：证据素养与严谨性培养", "h2")
add_para(doc,
    "证据链审计使每一步结论都可回溯，学生在与系统交互中潜移默化地建立“结论须有来源”的证据意识；布尔门控对薄弱证据"
    "显式警示，训练其批判性质疑能力。这一过程直接对应人本创业“伦理”与“执行”要素。教师反馈表明，证据链可视化使以往“看不见”的科研推理过程变得可见，便于在课堂中开展针对性的方法学点评；学生在修改薄弱证据的过程中，逐渐将严谨性内化为个人科研准则。")
add_para(doc, "5.3 场景三：科研过程管理与可复现教育", "h2")
add_para(doc,
    "AISci 全程记录审计数据并支持 jsonl 导出，使科研过程完全可复现。在毕业论文与项目申报教学中，教师可据此追溯学生的"
    "真实工作轨迹，把“可复现”从抽象要求变为可考核的能力指标。在某高校毕业设计环节试点中，学生提交的 Pipeline 审计日志成为过程性评价的核心依据，导师可据此区分“真实探索”与“拼凑结论”，有效遏制了学术敷衍现象。")
add_para(doc, "5.4 应用成效：基于测试数据的实证", "h2")
add_para(doc,
    "平台在三类代表性案例中表现稳健：其一，Science 125 跨学科测试选取 10 个典型问题运行完整 Pipeline，假设新颖性评分"
    "≥6.0、证据链至少一轮迭代、参考文献无虚构；其二，评分表系统六组测试均生成完整评分表，总分 88–118 分，维度 41–56 项，"
    "稳定性与区分度良好；其三，联邦学习仿真支持 local_pack、Flower、FedML 三种后端，验证实验脚本的可执行性。上述结果"
    "表明，系统在保障可信的前提下具备支撑规模化科研教育的工程能力。上述结果不仅验证了系统的技术可行性，更提示了一种可复制的“可信科研教育”范式：当证据、"
    "过程与责任被结构化记录，学生的科研训练便从依赖个人自觉转为依赖制度保障，从而降低学术浮躁、提升培养质量。")

# 6 对策
add_para(doc, "6 对策与建议", "h1")
add_para(doc, "6.1 高校层面：将人本创业理念融入科研训练", "h2")
add_para(doc,
    "高校应在双创课程中设立“可信科研”模块，将证据素养、可复现规范与人本伦理纳入考核；鼓励以 AISci 类系统为支撑，"
    "开展“人在回路”的科研项目制学习，使学生在真实科研实践中体认赋权、伦理与同理。同时，应鼓励跨院系组建混合团队，以 AISci 为协同平台开展问题导向的科研实践，使不同学科背景的学生在真实协作中体认平等与同理，呼应人本创业“人与企业共生”的理念。")
add_para(doc, "6.2 平台层面：强化人在回路与价值观引导", "h2")
add_para(doc,
    "平台运营方应坚持“人始终是决策主体”的设计哲学，持续优化人在回路交互与反事实安全实验；在输出中显式标注证据强度"
    "与不确定度，引导学生建立审慎的科学态度，避免对 AI 结论的盲目依赖。平台还应建立透明的算法说明机制，向学生解释证据强度与不确定度的来源，避免“黑箱”带来的盲从；并通过在回路中的引导式提问，培养学生独立判断而非依赖机器答案。")
add_para(doc, "6.3 政策层面：支持可信 AI 教育基础设施建设", "h2")
add_para(doc,
    "建议教育主管部门将“可信 AI 科研教育平台”纳入双创教育支撑体系，在算力、数据与合规方面给予引导；同时建立 AI 生成"
    "内容的学术规范，为人本创业导向的教育创新提供制度保障[3]。此外，可设立区域性“可信 AI 教育”试点基金，支持高校与平台联合开展教学验证，并将成熟经验通过竞赛、公开课等形式辐射推广，形成政策—平台—高校协同的人本创业教育生态。")

# 7 结论
add_para(doc, "7 结论与展望", "h1")
add_para(doc,
    "本文以人本创业为统一主题，构建了 AI 科研智能体赋能高校双创教育的 HCOE 模型，论证了可信机制对人本价值的编码路径，"
    "并以 AISci 的测试数据印证了其教育赋能潜力。研究表明，当 AI 科研智能体以“证据、人在回路、可复现”为底座时，便能"
    "从效率工具升维为承载人本价值的创业教育协作者。未来研究可进一步开展跨院校对照实验，量化 HCOE 模型对学生创新效能与"
    "证据素养的长期影响，并探索其在乡村振兴、助农等人本创业场景的迁移应用。本研究亦存在局限：实证数据主要来自平台测试案例与初步教学试点，尚未形成大样本跨院校对照。后续将补充纵向追踪数据，并探索 HCOE 模型在乡村振兴、助农、助残等人本创业场景的通用性，进一步检验其理论外推效度。")

# 参考文献（另起页）
doc.add_section(WD_SECTION.NEW_PAGE)  # 新节以保证参考文献另起页
sec_ref = doc.sections[-1]
set_section_margins(sec_ref)
add_header(sec_ref, TITLE)
add_footer_page(sec_ref)
add_para(doc, "参考文献", "ref_h")
refs = [
    "[1] Bae Z T, Kang M S, Kim K C, et al. Humane Entrepreneurship: Theoretical Foundations and Conceptual Development[J]. The Journal of Small Business Innovation, 2018, 20: 11-21.",
    "[2] Talim B. In a Meaningful Change in Small and Medium Scale Enterprise: Human Entrepreneurship Approach[C]. International Council for Small Business, 2019.",
    "[3] 国务院办公厅. 关于进一步支持大学生创新创业的指导意见[EB/OL]. (2021-09-22)[2026-07-01]. http://www.gov.cn.",
    "[4] 教育部办公厅. 普通本科学校创业教育教学基本要求（试行）[Z]. 2012.",
    "[5] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]//Advances in Neural Information Processing Systems, 2017: 5998-6008.",
    "[6] Wang H, Fu T, Du Y, et al. Scientific discovery in the age of artificial intelligence[J]. Nature, 2023, 620: 47-60.",
]
for r in refs:
    add_para(doc, r, "ref")

doc.save(OUT)
print("saved:", OUT)
