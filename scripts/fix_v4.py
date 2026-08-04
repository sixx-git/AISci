# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PATH = "D:/Workplace/AISci/output/联邦智研_大创申报书_v4.docx"
doc = Document(PATH)


def set_run_font(run, cn='宋体', en='宋体', size=None, bold=None):
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), cn)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)


def new_para(ref, text, before=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             first_line=None, left=None, ls=1.5, sa=Pt(4), bold=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    pf = p.paragraph_format
    pf.alignment = align
    if first_line is not None:
        pf.first_line_indent = first_line
    if left is not None:
        pf.left_indent = left
    pf.line_spacing = ls
    pf.space_after = sa
    r = p.add_run(text)
    set_run_font(r, cn='宋体', en='宋体', bold=bold)
    if before:
        ref._p.addprevious(p._p)
    else:
        ref._p.addnext(p._p)
    return p


# ---------- FIX 1: restore (1)证据溯源 + shorten the real 1.1 long sentence ----------
p_wrong = None   # the one I wrongly overwrote (my shortened "提出系统")
p_long = None    # the original 1.1 "提出系统" long sentence
for p in doc.paragraphs:
    t = p.text
    if '提出一种基于国产大模型 Qwen' in t:        # my earlier (wrong) overwrite, has space before Qwen
        p_wrong = p
    elif '提出了一种基于国产大模型Qwen' in t:      # original 1.1 long sentence, no space before Qwen
        p_long = p

assert p_wrong is not None, "p_wrong not found"
assert p_long is not None, "p_long not found"

# restore the (1)证据溯源问题 content (was destroyed)
p_wrong.runs[0].text = ('（1）证据溯源问题：如何确保系统生成的每一条科学假设都能追溯到真实的文献来源，'
                        '杜绝虚构引用？解决方案：设计 Fact 白名单约束机制，LLM 在生成和修订假设时仅允许引用'
                        '经过验证的文献事实（fact_id），并自动过滤未授权的引用。')

# shorten the real 1.1 long sentence (keeps the "提出系统" framing, removes near-duplicate)
p_long.runs[0].text = ('针对上述挑战，本项目“联邦智研”（AISci）提出一种基于国产大模型 Qwen 的多智能体科研自动化系统，'
                       '以“矛盾→拆解→空白→假设→验证”的科学逻辑链为核心，构建从研究问题到可验证假设的全流程自动化 Pipeline，'
                       '并通过证据链迭代推理、布尔质量门禁、反事实预演等机制，系统性解决 LLM 在科研中的幻觉引用与质量控制问题。')

# ---------- FIX 2: populate the static 目录 with a manual TOC ----------
toc_head = None
for p in doc.paragraphs:
    if p.text.strip() == '目  录':
        toc_head = p
        break
assert toc_head is not None, "目录 heading not found"

toc_entries = [
    "摘要",
    "核心创新点",
    "一、研究问题与解决方法",
    "二、架构设计与讲解",
    "三、项目工作流程",
    "四、上下文工程设计",
    "五、数据或资料来源说明",
    "六、代表性测试案例",
    "七、结果展示与反馈迭代过程",
    "八、应用示范：科学影响力预测系统",
    "九、源代码与可复现性",
    "十、总结与展望",
    "附录",
    "附录二：科学影响力预测系统技术细节",
]

# chain-insert after the 目录 heading so order is preserved
anchor = toc_head
for e in toc_entries:
    anchor = new_para(anchor, e, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=None, left=Pt(12), sa=Pt(2))

doc.save(PATH)
print("FIX APPLIED. Restored 证据溯源:", p_wrong is not None, "| Shortened 1.1:", p_long is not None,
      "| TOC entries:", len(toc_entries))
