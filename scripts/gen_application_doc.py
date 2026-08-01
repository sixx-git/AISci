"""
生成"联邦智研"大创申报书 Word 文档
遵循 competition-doc-standards 规范（计划书/方案类格式）
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set cell background color — removes existing shading first to avoid duplicates."""
    from lxml import etree
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>')
    tcPr.append(shading_elm)

def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a formatted table — 表头宋体加粗 灰底 / 数据宋体."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        rPr = run._element.get_or_add_rPr()
        rFonts_elem = rPr.find(qn('w:rFonts'))
        if rFonts_elem is None:
            from lxml import etree as _etree
            rFonts_elem = _etree.SubElement(rPr, qn('w:rFonts'))
        rFonts_elem.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'CFCDCD')

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10.5)
            run.font.name = 'Times New Roman'
            rPr = run._element.get_or_add_rPr()
            rFonts_elem = rPr.find(qn('w:rFonts'))
            if rFonts_elem is None:
                from lxml import etree as _etree2
                rFonts_elem = _etree2.SubElement(rPr, qn('w:rFonts'))
            rFonts_elem.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()
    return table

def add_heading1(doc, text):
    """Add Chapter heading (Heading 1) — 黑体 bold 小三号 居中."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree
        rFonts_elem = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading2(doc, text):
    """Add Section heading (Heading 2) — 黑体 bold 四号 左对齐."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree
        rFonts_elem = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading3(doc, text):
    """Add Sub-section heading (Heading 3) — 黑体 bold 小四 左对齐."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree
        rFonts_elem = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_body(doc, text):
    """Add body paragraph — 宋体 小四 1.5倍行距 首行缩进."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree
        rFonts_elem = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_body_no_indent(doc, text):
    """Add body paragraph without indent — 宋体 小四 1.5倍行距."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree
        rFonts_elem = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_code_block(doc, code_text):
    """Add a code block paragraph — Consolas 小五号 左对齐."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    rPr = run._element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree
        rFonts_elem = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(50, 50, 50)

def main():
    doc = Document()

    # ===== Page Setup =====
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ===== Set default font =====
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree as _e
        rFonts_elem = _e.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ===== COVER PAGE =====
    # Add spacing before the cover title — use fewer paragraphs with proper spacing
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    # Main title on cover — 黑体 初号 居中对齐
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('联邦智研')
    run.bold = True
    run.font.size = Pt(42)
    run.font.name = '黑体'
    rPr = run._element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        from lxml import etree as _e2
        rFonts_elem = _e2.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after = Pt(24)
    run2 = p2.add_run('——基于Qwen的多智能体科研自动化系统')
    run2.font.size = Pt(16)
    run2.font.name = 'Times New Roman'
    rPr2 = run2._element.get_or_add_rPr()
    rFonts_elem2 = rPr2.find(qn('w:rFonts'))
    if rFonts_elem2 is None:
        rFonts_elem2 = _e2.SubElement(rPr2, qn('w:rFonts'))
    rFonts_elem2.set(qn('w:eastAsia'), '宋体')
    run2.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle spacing
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    p3.paragraph_format.space_after = Pt(18)
    run3 = p3.add_run('大学生创新创业大赛申报书')
    run3.font.size = Pt(22)
    run3.font.name = 'Times New Roman'
    rPr3 = run3._element.get_or_add_rPr()
    rFonts_elem3 = rPr3.find(qn('w:rFonts'))
    if rFonts_elem3 is None:
        rFonts_elem3 = _e2.SubElement(rPr3, qn('w:rFonts'))
    rFonts_elem3.set(qn('w:eastAsia'), '黑体')
    run3.bold = True
    run3.font.color.rgb = RGBColor(0, 0, 0)

    # Cover info table
    cover_info = [
        ('参赛赛道', '高教主赛道·人工智能+'),
        ('参赛组别', '创意组'),
        ('项目类型', '人工智能'),
        ('作品方向', 'A. 科学假设生成与研究计划设计'),
        ('技术底座', 'Qwen（千问）/ 阿里云百炼'),
    ]
    table = doc.add_table(rows=len(cover_info), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(cover_info):
        cell_label = table.rows[i].cells[0]
        cell_label.text = ''
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        rPr_l = run._element.get_or_add_rPr()
        rFonts_l = rPr_l.find(qn('w:rFonts'))
        if rFonts_l is None:
            rFonts_l = _e2.SubElement(rPr_l, qn('w:rFonts'))
        rFonts_l.set(qn('w:eastAsia'), '黑体')

        cell_value = table.rows[i].cells[1]
        cell_value.text = ''
        p = cell_value.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        rPr_v = run._element.get_or_add_rPr()
        rFonts_v = rPr_v.find(qn('w:rFonts'))
        if rFonts_v is None:
            rFonts_v = _e2.SubElement(rPr_v, qn('w:rFonts'))
        rFonts_v.set(qn('w:eastAsia'), '宋体')

    # Page break after cover
    doc.add_page_break()

    # ===== TABLE OF CONTENTS placeholder =====
    add_heading1(doc, '目  录')
    add_body_no_indent(doc, '（自动目录，请在Word中插入：引用 → 目录 → 自动目录）')
    doc.add_page_break()

    # ===== CHAPTER 1 =====
    add_heading1(doc, '一、研究问题与解决方法')

    add_heading2(doc, '1.1 研究背景与科学问题')
    add_body(doc, '科学研究正面临信息爆炸与认知瓶颈的双重挑战。据统计，全球每年发表超过300万篇学术论文，科研人员难以全面追踪本领域的前沿进展，导致大量重复性研究和被忽视的知识交叉点。以《Science》发布的125个前沿科学问题为例，这些问题横跨数学、物理、化学、生物、医学等多个学科，传统的人工文献检索与假设提出方式效率低下，难以系统性地发现跨学科知识缺口。')
    add_body(doc, '近年来，大语言模型（LLM）在文本理解、知识推理和内容生成方面展现出强大能力。然而，单纯依赖LLM进行科学假设生成面临三大核心痛点：一是LLM容易产生"幻觉引用"（Hallucination），生成不存在的文献或事实；二是缺少结构化的证据链推理，假设缺乏可追溯性和可验证性；三是缺乏系统化的质量控制机制，无法自动判断假设的新颖性、可行性和与问题的匹配度。')
    add_body(doc, '针对上述挑战，本项目"联邦智研"（AISci）提出了一种基于国产大模型Qwen的多智能体科研自动化系统。系统以"矛盾→主要矛盾→拆解→研究对象→研究现状→知识空白→假设→方法"的科学逻辑链为核心，构建了从研究问题到可验证科学假设的全流程自动化Pipeline，并通过证据链迭代推理、布尔质量门禁、反事实预演等机制有效解决了LLM在科研场景中的幻觉和质量控制问题。')

    add_heading2(doc, '1.2 核心研究问题')
    add_body(doc, '本项目的核心研究问题聚焦于：如何构建一个人机协作的多智能体系统，使其能够基于真实文献证据自动生成可验证的科学假设，并具备闭环自迭代优化能力？')
    add_body(doc, '具体而言，本项目致力于解决以下三个子问题：')

    add_body(doc, '（1）证据溯源问题：如何确保系统生成的每一条科学假设都能追溯到真实的文献来源，杜绝虚构引用？解决方案：设计Fact白名单约束机制，LLM在生成和修订假设时仅允许引用经过验证的文献事实（fact_id），并自动过滤未授权的引用。')
    add_body(doc, '（2）质量控制问题：如何系统化地评估和筛选科学假设的质量？解决方案：构建布尔质量门禁系统（11种阶段特定Gate），替代传统的连续评分方式，在每个关键节点（新颖性审查、Ensemble评审、沙箱验证、证据链完备性等）进行离散化质量判定。')
    add_body(doc, '（3）迭代优化问题：如何让系统根据反馈（包括实验结果、评审意见、文献补充）自动优化假设和实验方案？解决方案：设计科学自迭代编排器，通过"证据弱→补文献→重跑Pipeline"的闭环机制，实现多轮精化迭代。')

    add_heading2(doc, '1.3 技术路线与解决思路')
    add_body(doc, '本项目的技术路线如图1-1所示，核心思路是将科学研究的逻辑链（矛盾→拆解→空白→假设→验证）编码为多智能体协作Pipeline，并以证据链迭代推理引擎和布尔质量门禁系统为两大支柱，确保系统的可靠性和产出质量。')

    add_body(doc, '图1-1：AISci系统功能示意图（见附录）')

    # ===== CHAPTER 2 =====
    add_heading1(doc, '二、架构设计与讲解')

    add_heading2(doc, '2.1 系统整体架构')
    add_body(doc, '本项目采用前后端分离的三层架构：后端基于Python FastAPI构建，提供18个API路由模块；前端基于React 18 + Vite 5 + TailwindCSS 3构建，采用Blueprint深色设计风格；数据层使用SQLite数据库配合zvec向量检索，支持文献全文的语义搜索和相似度匹配。')

    add_body(doc, '系统整体架构遵循"Agent-Skill-Infrastructure"三层分离原则：')

    add_table_with_style(doc,
        ['架构层', '技术选型', '核心职责'],
        [
            ['智能体层 (Agent)', '6个独立Agent类 + Prompt模板', '流程编排与LLM调用'],
            ['技能层 (Skill)', '约70个Skill模块/8个子领域', '文献检索、证据推理、数据清洗等可复用工具'],
            ['基础设施层 (Infra)', 'FastAPI + zvec + SQLite + PyMuPDF', 'API服务、向量检索、数据库、PDF解析'],
        ])

    add_heading2(doc, '2.2 多智能体协作Pipeline设计')
    add_body(doc, '系统的核心是七阶段顺序Pipeline，每个阶段由独立的Agent负责，通过标准化的JSON Schema确保阶段间数据传递的类型安全。Pipeline的阶段顺序为：')

    add_body(doc, '问题理解（Problem Understanding）→ 文献挖掘（Literature Mining）→ 知识缺口发现（Knowledge Gap）→ 假设生成（Hypothesis Generation）→ 假设评审（Hypothesis Review）→ 迭代实验（Iterative Experiment）→ 报告生成（Report Generation）')

    add_body(doc, '每个阶段执行后自动记录完整审计数据：输入数据（input_data）、输出数据（output_data）、使用的Prompt原文（prompt_used）、模型参数（model_parameters）、Token用量（token_count）、耗时（duration_ms）以及每次Qwen API调用的独立记录（CallLog）。全链路审计数据支持jsonl格式导出，确保科研过程的完全可复现。')

    add_heading2(doc, '2.3 核心智能体设计')
    add_body(doc, '系统包含6个核心智能体，每个智能体均采用"独立类 + Prompt模板 + Pydantic JSON Schema"的设计模式：')

    add_table_with_style(doc,
        ['智能体', '输入', '输出', '核心职责'],
        [
            ['ProblemUnderstanding', '研究问题文本', '结构化问题描述（含主要矛盾、研究对象、边界）', '将模糊问题转化为可研究的结构化描述'],
            ['LiteratureMining', 'project_id + 研究问题', '科学事实列表 + 证据列表 + 引用映射', '向量检索→LLM提取事实，每条绑定chunk_id'],
            ['KnowledgeGap', '文献事实 + 不确定点', '知识缺口 + 矛盾点 + 研究机会', '发现文献中的空白和可研究方向'],
            ['HypothesisGeneration', '知识缺口 + 文献证据', '候选假设 + 对齐评分 + 数据证据', '基于缺口生成可验证假设，标注supporting_fact_ids'],
            ['HypothesisReview', '候选假设 + 文献上下文', '新颖性评分 + 可行性评估 + 反事实预演', '评估假设质量，执行L0 CounterfactualPreview'],
            ['ReportGeneration', '全流程中间产物', '12字段最终报告 + 合规检查', '聚合各阶段输出，生成完整报告'],
        ])

    add_heading2(doc, '2.4 技能工具层（Skill Layer）设计')
    add_body(doc, 'Skill层是Agent层的可复用工具库，按科研子领域组织为8大类别、约70个工具模块。Skill与Agent的关系是：Agent负责流程编排和LLM调用，Skill提供具体的研究工具能力（如arXiv文献搜索、PDF表格抽取、证据链迭代等）。')

    add_table_with_style(doc,
        ['Skill类别', '代表模块', '核心功能'],
        [
            ['文献 (Literature)', 'arxiv_search_skill, citation_grounding_skill', '多源检索、引用真实性验证'],
            ['Data Finder', 'pdf_table_extraction_skill, dataset_merge_skill', 'PDF表格抽取、Schema对齐、实体对齐'],
            ['证据推理 (Evidence)', 'iterative_hypothesis_loop_skill, hypothesis_revision_skill', '多轮证据检索、LLM假设修订（Fact白名单）'],
            ['多模态 (Multimodal)', 'qwen_vl_image_understanding_skill', 'VLM图像理解、多模态fact构建'],
            ['推理 (Reasoning)', 'hypothesis_novelty_review_skill, hypothesis_tournament_skill', '新颖性审查、锦标赛假设选择'],
            ['报告 (Report)', 'scientific_plot_skill, report_quality_check_skill', '图表生成、12字段合规检查'],
            ['反事实 (Counterfactual)', 'counterfactual_preview_skill', 'FALSIFY过滤器、反事实预演'],
            ['实验 (Experiment)', 'experiment_sanity_check_skill', '实验方案合理性检查'],
        ])

    # ===== CHAPTER 3 =====
    add_heading1(doc, '三、项目工作流程')

    add_heading2(doc, '3.1 七阶段Pipeline完整流程')
    add_body(doc, '用户在前端创建研究项目并输入科学问题后，系统通过POST /api/v1/pipeline/run接口触发Pipeline执行。整个流程从前端PipelineProgress组件实时展示进度，到RunLogDetail组件查看每步详情，实现了科研过程的全透明化和可追溯化。')

    add_body(doc, '阶段一：问题理解（Problem Understanding）——接收用户输入的研究问题文本，调用Qwen模型将模糊问题转化为结构化描述。核心产出包括：主要矛盾（main_contradiction）、现象矛盾来源（phenomenon_contradiction）、研究对象（research_object，含内/外/边界三维拆解）、研究目的（expected_output）等字段。')

    add_body(doc, '阶段二：文献挖掘（Literature Mining）——基于研究问题和关键词，通过向量检索从文献库中召回相关文献片段，再由Qwen模型从每个片段中提取结构化科学事实。每条事实绑定来源chunk_id，确保可追溯。同时支持arXiv在线检索和PDF全文解析。')

    add_body(doc, '阶段三：知识缺口发现（Knowledge Gap）——基于已提取的文献事实和不确定点，由LLM分析现有研究的覆盖范围，识别尚未被充分研究的空白区域（knowledge_gaps），标注潜在的矛盾点和可突破的研究机会。')

    add_body(doc, '阶段四：假设生成（Hypothesis Generation）——基于知识缺口和文献证据，LLM生成多个候选假设。每个假设都标注了supporting_fact_ids，确保与文献证据的绑定。同时检查假设是否偏题（通过领域关键词检测），并通过Pairwise Tournament进行假设排序。')

    add_body(doc, '阶段五：假设评审（Hypothesis Review）——对候选假设进行多维评估：新颖性评分（是否为已有工作的重复）、可行性评估（实验方案是否可执行）、反事实预演（如果假设不成立会怎样）。评审结果通过Ensemble Gate判定是否通过。')

    add_body(doc, '阶段六：迭代实验（Iterative Experiment）——绑定数据集→设计分析脚本→沙箱执行→结果分析→脚本重设计，形成Plan→Execute→Analyze→Reflect的闭环。实验计划需通过可执行性Gate（数据列匹配验证）。')

    add_body(doc, '阶段七：报告生成（Report Generation）——聚合全流程中间产物，生成包含12个标准字段（Paper Title、Abstract、Problem Statement、Rationale、Technical Details、Datasets、Source、Target、Methods、Experiments、Results、References）的完整科研报告，并通过ReportQualityCheckSkill进行合规检查。')

    add_heading2(doc, '3.2 HITL人机协作机制')
    add_body(doc, '系统在关键Pipeline节点设置了HITL（Human-in-the-Loop）人工审核Gate。当布尔门禁判定某一阶段的质量未达标时，Pipeline自动暂停等待人工审核。用户可以通过前端界面：①查看和编辑阶段输出（save_stage_human_output）；②从任意阶段重跑Pipeline（rerun_from_stage）；③与阶段输出进行交互式对话修改（stage_chat）；④获取导师级反馈（mentor_review）。人工审核通过后，Pipeline继续执行。')

    add_heading2(doc, '3.3 科学自迭代闭环')
    add_body(doc, '当假设评审未通过（auto_triggers=review_reject）或证据强度不足（auto_triggers=evidence_weak）时，ScienceIterationOrchestrator自动触发自迭代编排：在证据弱时自动补充文献→重新构建证据链→重新运行假设树→再次评审。迭代过程通过closed_loop_decisions模块记录完整的因果链摘要，自动推断每次迭代的驱动来源（validation_feedback/literature_refresh/human_feedback/ensemble_review），最终通过VersionComparePanel展示迭代前后的质量对比。')

    # ===== CHAPTER 4 =====
    add_heading1(doc, '四、上下文工程设计')

    add_heading2(doc, '4.1 多轮证据链迭代推理引擎（核心创新一）')
    add_body(doc, '证据链迭代推理引擎是本项目最具学术价值的算法创新。它构建了一个8步骤串联的证据-假设对弈管道：科学声明提取（ScientificClaimExtraction）→ 支持证据检索（EvidenceRetrieval）→ 反证据检索（CounterEvidenceRetrieval）→ 证据立场分类（EvidenceStanceClassification）→ 假设修订（HypothesisRevision）→ 证据接地（EvidenceGrounding）→ 引用完整性检查（CitationIntegrityCheck）→ 证据链构建（EvidenceChainBuilder）。')
    add_body(doc, '该引擎的核心创新包括：')
    add_body(doc, '（1）Fact白名单强制约束：LLM在修订假设时，只能引用fact_whitelist中的fact_id。修订后还会对cited_fact_ids做白名单过滤验证，从机制层面杜绝LLM幻觉引用。这一设计解决了当前LLM在科学写作中最被诟病的"编造参考文献"问题。')
    add_body(doc, '（2）证据收敛终止条件：每轮迭代都重新评估支持证据和反对证据的存在情况。达到最大迭代轮次或"存在支持证据且无反对证据"时自动终止，避免过度迭代导致的假设退化。')
    add_body(doc, '（3）证据立场三级分类策略：优先采用检索阶段已标记的stance（support/refute/neutral）；其次通过COUNTER_KEYWORDS（如"limitation"、"failure"、"限制"、"失败"等14个关键词）检测反对立场；最后通过假设核心术语在证据中的出现情况判定支持/中性立场。')
    add_body(doc, '（4）中英混合分词相关性评分：对英文按词切分，对中文按二元组（bigram）切分，解决了CJK文本无空格导致的Jaccard相关度恒为0的问题。最终使用max(jaccard×1.4, coverage×0.9)的混合评分策略。')

    add_heading2(doc, '4.2 布尔质量门禁系统（核心创新二）')
    add_body(doc, '传统科研辅助工具多采用0-100连续评分来评估输出质量，但连续评分在科研场景中缺乏明确的决策支持——70分是否意味着"可以通过"？60分和65分之间的差异在科研判断上有何实质含义？')
    add_body(doc, '本系统设计了11种阶段特定的布尔Gate：新颖性Gate（gate_novelty，新颖性≥6.0且非高风险重叠）、Ensemble Gate（gate_ensemble，评审≥6.5或Accept）、沙箱Gate（gate_sandbox，沙箱实测成功）、联邦Gate（gate_federated，联邦双门槛）、图表Gate（gate_plot，图表≥6.5）、覆盖度Gate（gate_coverage，覆盖度≥70%）、证据链Gate（gate_evidence，至少1轮证据链迭代）、HITL Gate（gate_hitl，人工审核已处理）、可执行性Gate、验收Gate（gate_acceptance）、CoT Gate。')
    add_body(doc, '每个Gate输出明确的PASS/FAIL判断。趋势追踪函数（summarize_gate_trend）跟踪连续失败次数和改善趋势，支持停滞停止决策——当连续多轮未通过且无明显改善趋势时，自动触发HITL暂停。')

    add_heading2(doc, '4.3 统一反馈中心（核心创新三）')
    add_body(doc, '在多智能体协作系统中，如何将某一阶段的发现有效传递给后续阶段是一个关键挑战。传统方法通常依赖硬编码的阶段间数据传递，缺乏灵活性和可扩展性。')
    add_body(doc, '本项目设计了统一反馈中心（Feedback Hub），汇聚来自7种来源的反馈信息：HITL人工审核、知识图谱（KG）、数据查找器（Data Finder）、溯源追踪（Provenance）、文献检索（Literature）、用户直接输入（User）和多模态分析（Multimodal）。所有反馈被转换为约束文本，存入global_constraints池（最多保留50条），在Pipeline后续阶段的执行中被get_active_constraints自动拉取并注入到LLM的上下文中。')
    add_body(doc, '反馈中心还维护了STAGE_TO_FEEDBACK_TARGET映射表和RERUN_TARGETS定义，使得不同类型的反馈可以精确触发对应阶段的重跑操作。')

    add_heading2(doc, '4.4 反事实预演机制（核心创新四）')
    add_body(doc, '在科学假设评审和实验设计之间，系统插入了一个L0级别的定性反事实预演层（CounterfactualPreview）。预演的核心是FALSIFY过滤器（filter_falsify_scenarios）：保留可证伪（falsifiable=True）、有证据支撑（evidence_fact_ids在有效集合内）、能指导实验（有cheap_test方案）的反事实场景，并过滤掉不可证伪或缺乏决策影响（无decision_impact）的场景。')
    add_body(doc, '通过build_counterfactual_feedback_constraints函数，预演结果被转换为实验设计阶段的全局约束，特别是高风险场景的对照组构建建议。当存在高风险反事实场景且无失败预测时，proceed_to_iterative_experiment被设为False，阻断实验流程，保护科研资源不被浪费在无价值的实验设计上。')

    add_heading2(doc, '4.5 实验计划可执行性评估')
    add_body(doc, '针对LLM生成的实验计划"看起来合理但实际无法执行"的问题，系统设计了实验计划可执行性Gate。其工作原理是：①从data_context和实验设计的verifiable_hypothesis/methods/experimental_steps中提取需求信号；②用正则模式从信号文本中推断所需数据列名；③与可用数据列做精确匹配+模糊子串匹配；④按评分公式计算可执行性得分：score = 40 + (has_data?25:0) + (has_steps?15:0) + (has_metrics?10:0) + min(10, matched×2) − min(30, blockers×10)。通过条件为score≥60、无blockers、有steps、有metrics。')

    add_heading2(doc, '4.6 Pairwise锦标赛假设选择')
    add_body(doc, '参考Sakana AI Scientist（arxiv:2408.06292）的锦标赛假设选择机制，本项目实现了改进版的Margin-Weighted Tournament。O(n²)全配对比较中，每对由LLM作为裁判进行4维评估（新颖性、可验证性、与证据一致性、可行性）。胜者得1.0 + margin分，败者得margin×0.3分，平局各得0.5分。margin来自LLM输出的置信度，比原始Tournament仅做胜负二值判定更精细。')

    # ===== CHAPTER 5 =====
    add_heading1(doc, '五、数据或资料来源说明')

    add_heading2(doc, '5.1 文献数据源')
    add_body(doc, '系统的文献挖掘阶段支持多种数据来源：')
    add_body(doc, '（1）arXiv开放学术论文库：通过arxiv_search_skill进行关键词检索，自动下载并解析PDF全文（PyMuPDF + pdfplumber双引擎），提取结构化科学事实。')
    add_body(doc, '（2）HuggingFace Datasets：支持通过data_finder服务检索和导入数据集，可自动导入样例行数据。')
    add_body(doc, '（3）Zenodo / Figshare开放数据存储库：支持检索和自动导入（≤25MB文件），获取科研数据集的元数据和样例行。')
    add_body(doc, '（4）Kaggle / OpenAlex / NCBI GEO：提供元数据检索和静态索引查询，辅助数据集发现。')
    add_body(doc, '（5）用户上传：支持PDF、CSV、TXT等格式的文献和数据文件直接上传到项目文献库。')

    add_heading2(doc, '5.2 科学问题基准数据集')
    add_body(doc, '系统使用《Science》杂志发布的"125个前沿科学问题"作为核心测试基准。项目已将125个问题整理为英中平行数据集（output/sjtu-125-questions/），涵盖数学科学（3题）、化学（7题）、物理（7题）、生物学（14题）、医学（18题）、神经科学（9题）等14个学科分类，每题附带context背景段落，用于验证系统在跨学科科学假设生成中的通用性。')

    add_heading2(doc, '5.3 联邦学习参考数据')
    add_body(doc, '联邦学习（FL）模式下，系统内置了丰富的参考数据资源：')
    add_body(doc, '（1）14个数据集元数据配置（YAML格式）：涵盖LEAF FEMNIST（手写字符识别）、LEAF Shakespeare（文本联邦学习）、UCI Adult（垂直联邦学习两方拆分）、PEMS Traffic（交通流预测）、SISFall HAR（人体活动识别）等经典FL benchmark。')
    add_body(doc, '（2）31篇FL领域核心论文的结构化事实（facts.json）：包括FedAvg（mcmahan2017）、FedProx（li2020）、LEAF（caldas2018）、SplitNN（vepakomma2018）、差分隐私（abadi2016）、安全聚合（bonawitz2017）、LoRA（hu2022）、OpenFedLLM（ye2024）等，覆盖联邦学习的所有核心技术方向。')
    add_body(doc, '（3）15个领域标签的预解析文献种子：包括金融风控、医疗健康、边缘计算、工业互联网、智慧交通、差分隐私、联邦CV、联邦NLP、联邦多语言、联邦LLM微调、LoRA异构、区块链、联邦强化学习、持续学习等。')

    # ===== CHAPTER 6 =====
    add_heading1(doc, '六、代表性测试案例')

    add_heading2(doc, '6.1 案例一：基于Science 125问题的科学假设生成')
    add_body(doc, '测试目标：验证系统能否基于《Science》125个前沿科学问题自动生成有证据支撑、可验证的科学假设。')
    add_body(doc, '测试方法：从125个问题中选取不同学科的10个典型问题作为输入，运行完整7阶段Pipeline，评估输出假设的新颖性、与文献的一致性、可验证性和报告完整性。')
    add_body(doc, '预期验证指标：假设新颖性评分≥6.0（gate_novelty通过标准）；证据链至少1轮迭代（gate_evidence通过标准）；报告中References字段无虚构引用（quality_check.critical_issues为空）。')

    add_heading2(doc, '6.2 案例二：评分表系统的科学影响力评估')
    add_body(doc, '测试目标：验证pingfenbiao评分表系统能否对学术报告进行准确的科学影响力评估。')
    add_body(doc, '测试方法：使用3个独立优化的生成器（Gen-1：主张核查、Gen-2：数据分析、Gen-3：科学调研）在FL（联邦学习）和PEFT（参数高效微调）两个领域进行了6组测试。')
    add_body(doc, '测试结果概要：6组测试均成功生成完整的评分表（task.json）和评分结果（rubric_scores.json），总评分范围88-118分，评分维度数41-56项，表明评分系统具有较好的稳定性和区分度。')

    add_heading2(doc, '6.3 案例三：联邦学习仿真测试')
    add_body(doc, '测试目标：验证FL Starter Pack能否正确注入联邦学习内容，驱动实验设计和仿真。')
    add_body(doc, '测试方法：创建联邦学习模式项目→选择标准Non-IID实验范式（Dirichlet α=0.1 + FedAvg/FedProx）→运行Pipeline→在迭代实验阶段验证仿真脚本的可执行性。支持local_pack、Flower（可选flwr）和FedML（可选fedml）三种仿真后端。')

    add_heading2(doc, '6.4 案例四：端到端验收测试')
    add_body(doc, '测试目标：验证系统的接口完整性和基本功能可用性。')
    add_body(doc, '测试方法：运行scripts/check_e2e.py脚本，对14项核心检查点进行自动化验收。')
    add_body(doc, '测试结果：14项检查全部通过（14 PASS，0 WARN，0 FAIL），涵盖后端健康检查、LLM客户端诊断、项目CRUD、文献源列表、arXiv搜索、Pipeline运行状态、报告接口、假设Agent接口、Skills完整性等。')

    # ===== CHAPTER 7 =====
    add_heading1(doc, '七、结果展示与反馈迭代过程')

    add_heading2(doc, '7.1 假设生成质量保障体系')
    add_body(doc, '系统的假设质量保障通过四层机制实现：')
    add_body(doc, '第一层——领域对齐检测：hypothesis_generation_agent内置6个领域关键词集合（医学神经、医学肿瘤、医学临床、分子生物、社会政策、心理学），在生成假设后自动检测假设关键词是否属于研究问题的领域范围（_QUESTION_DOMAIN_KEYWORDS映射），过滤偏题假设（OFF_DOMAIN_KEYWORDS）。')
    add_body(doc, '第二层——新颖性审查：hypothesis_novelty_review_skill对每一条候选假设进行新颖性评分，检测与已有文献的高风险重叠，输出1-10分的新颖性评分和具体的重叠风险等级。')
    add_body(doc, '第三层——锦标赛排序：通过改进的Margin-Weighted Tournament对候选假设进行全配对比较，O(n²)交叉评估4个维度，输出排序后的假设列表和选择理由。')
    add_body(doc, '第四层——人工审核：HITL Gate在关键节点暂停，允许人工审核、编辑和决策。同时支持采用红蓝对抗后的演化候选（hypothesis_evolution，含simplify/out_of_box两种策略），但演化结果不自动覆盖主假设，需人工通过select-evolved-hypothesis API确认后才生效。')

    add_heading2(doc, '7.2 证据链完备性展示')
    add_body(doc, '系统通过Science Iteration机制确保证据链的完备性：当检测到证据强度不足（evidence_weak）时，auo_literature_on_weak_evidence自动触发补充文献检索→重新构建证据链→重新运行假设树→再次评审的单轮refine。前端证据链抽屉新增"来源"和"验证"Tab，调用provenance API展示假设溯源时间线（fact→多模态证据→数据集→spec的完整溯源链）。')
    add_body(doc, 'CloseLoopDecisions模块自动构建跨轮迭代的因果摘要，追踪数据变化（PDF表格新增数、CSV行数变化、完备性得分变化）、计划变化（实验步骤调整、研究方法调整、主假设文本修订、证据fact增量），并通过infer_driven_by函数自动归因迭代驱动来源。')

    add_heading2(doc, '7.3 迭代优化效果量化')
    add_body(doc, '系统的迭代优化通过以下机制量化展示：')
    add_body(doc, '（1）质量趋势追踪：summarize_gate_trend函数跟踪11种Gate的通过/失败历史，计算连续失败次数和改善趋势，支持可视化展示质量提升的演进过程。')
    add_body(doc, '（2）版本对比：VersionComparePanel联动IterationRoundPanel，展示每轮的评分delta（变化量）、资料补充计划和证据增量，直观呈现迭代带来的质量提升。')
    add_body(doc, '（3）审计链导出：完整审计链（quality_trend/events/decisions）支持jsonl格式导出，便于第三方审查和复现验证。')

    add_heading2(doc, '7.4 报告生成与质量检查')
    add_body(doc, 'ReportQualityCheckSkill对最终报告执行6项标准检查：12字段完整性（Paper Title→References）；Technical Details是否明确Qwen和阿里云百炼；References是否包含unknown/placeholder等虚构引用；Results是否区分actual/simulated/expected；Datasets是否有真实来源；图表是否有source_dataset_id和is_generated_from_real_data标记；是否出现非Qwen模型表述。检查结果以结构化JSON形式输出（score/passed/missing_fields/critical_issues/recommendations），前端QualityCheckCard组件可视化展示评分和建议。')

    # ===== CHAPTER 8 =====
    add_heading1(doc, '八、源代码与可复现性')

    add_heading2(doc, '8.1 项目代码结构')
    add_body(doc, '项目采用前后端分离架构，代码组织清晰：')
    add_code_block(doc, 'AISci/')
    add_code_block(doc, '├── backend/                    # Python后端 (FastAPI)')
    add_code_block(doc, '│   ├── app/agents/             # 6个智能体 (独立类+Prompt+Schema)')
    add_code_block(doc, '│   ├── app/api/                # 18个API路由模块')
    add_code_block(doc, '│   ├── app/core/               # 质量评分、闭环决策、溯源等核心模块')
    add_code_block(doc, '│   ├── app/services/           # 约55个业务服务 (Pipeline、科学迭代等)')
    add_code_block(doc, '│   ├── app/skills/             # 约70个技能模块 (8个子领域)')
    add_code_block(doc, '│   ├── app/integrations/       # 外部系统桥接 (shaxiang实验引擎)')
    add_code_block(doc, '│   ├── prompts/                # 阶段Prompt模板 + presets/范式预设库')
    add_code_block(doc, '│   ├── tests/                  # pytest测试套件 (7批A级优化回归)')
    add_code_block(doc, '│   └── data/                   # arXiv回退数据 + FL Starter Pack')
    add_code_block(doc, '├── frontend/                   # React前端 (Vite + TailwindCSS)')
    add_code_block(doc, '│   ├── src/pages/              # 8个页面 (首页、预测、文献等)')
    add_code_block(doc, '│   ├── src/components/         # 约60个组件 (工作流、假设卡片等)')
    add_code_block(doc, '│   └── src/services/           # API服务模块')
    add_code_block(doc, '├── pingfenbiao-main/           # 评分表系统 (独立Web服务)')
    add_code_block(doc, '└── shaxiang-main/              # 迭代实验引擎 (沙箱执行)')

    add_heading2(doc, '8.2 环境配置与启动')
    add_body(doc, '系统的环境配置简洁，仅需Python 3.10+和Node.js 18+。提供一键启动脚本：')
    add_code_block(doc, '# Windows一键启动')
    add_code_block(doc, 'scripts\\setup_backend.bat       # 创建venv + 安装依赖')
    add_code_block(doc, 'scripts\\setup_frontend.bat      # pnpm install')
    add_code_block(doc, 'scripts\\run_dev.bat             # 一键启动前后端')
    add_code_block(doc, '')
    add_code_block(doc, '# Linux/Mac一键启动')
    add_code_block(doc, 'bash scripts/setup_backend.sh')
    add_code_block(doc, 'bash scripts/setup_frontend.sh')
    add_code_block(doc, 'bash scripts/run_dev.sh')
    add_body(doc, '特别地，系统提供了USE_MOCK_LLM=true的Mock模式，无需真实Qwen API Key即可完整跑通Pipeline和前端交互，极大降低了评审专家的环境搭建成本。')

    add_heading2(doc, '8.3 测试覆盖与可复现性保证')
    add_body(doc, '项目的测试体系涵盖三个层次：')
    add_body(doc, '（1）单元测试：通过pytest对每个Agent类进行独立测试，验证输入输出Schema的正确性和边界条件处理。')
    add_body(doc, '（2）A级优化批次回归测试：覆盖1-7批A级优化功能（CQS+HITL、Verifiable Spec、Data Finder、闭环决策、文献图表、Feedback Hub、溯源审计），确保新功能不破坏已有能力。')
    add_body(doc, '（3）端到端验收测试：check_e2e.py脚本对14项核心检查点进行自动化验收（14 PASS，0 WARN，0 FAIL）。')
    add_body(doc, '审计链持久化路径（storage/audit/{run_id}.jsonl）确保每次Pipeline运行的完整记录可被复现和审查。')

    add_heading2(doc, '8.4 前端交互入口与交付完整度')
    add_body(doc, '系统提供完整的Web交互界面（React 18 + Vite 5 + TailwindCSS 3），涵盖12个功能Tab：项目概览、研究问题、文献库、智能体工作流、候选假设、迭代实验、研究报告、Prompt管理、运行日志、数据获取、Skills和设置。所有Tab通过23帧设计稿验收（Blueprint深色风格，深蓝底#0A1628 + 青色#38BDF8），确保UI一致性和规范性。')

    # ===== CHAPTER 9: 评选标准自评 =====
    add_heading1(doc, '九、评选标准自评对照')

    add_heading2(doc, '9.1 科学价值（40分）')
    add_table_with_style(doc,
        ['评分维度', '满分', '项目对应能力'],
        [
            ['科学事实表达准确性', '15', 'Fact白名单约束机制杜绝幻觉引用；每个fact绑定chunk_id可追溯；引用完整性自动检查'],
            ['内容转化/解释/展示清晰度', '15', '12字段标准报告+科学逻辑审查；前端图表可视化；开题逻辑三层次审查（Prompt/Schema/Skill）'],
            ['作品主题完整性与一致性', '10', '7阶段Pipeline覆盖完整科研流程；"矛盾→假设→验证"逻辑链贯穿始终'],
        ])

    add_heading2(doc, '9.2 技术深度（30分）')
    add_table_with_style(doc,
        ['评分维度', '满分', '项目对应能力'],
        [
            ['模型/智能体/技能设计完整性', '10', '6个Agent + 约70个Skill + 4套范式预设；Agent-Skill-Infra三层架构'],
            ['多模态数据处理/推理/交互能力', '10', 'Qwen-VL图像理解；PDF表格抽取；多模态证据构建；音频转录'],
            ['结果校验/反馈迭代与稳定性设计', '10', '11种布尔Gate；统一反馈中心；科学自迭代编排；闭环因果归因；停滞停止机制'],
        ])

    add_heading2(doc, '9.3 应用潜力（30分）')
    add_table_with_style(doc,
        ['评分维度', '满分', '项目对应能力'],
        [
            ['面向真实场景的使用价值', '10', '面向Science 125科学问题等真实研究场景；支持通用/联邦学习两种科研模式'],
            ['作品演示/交互入口与交付完整度', '10', '12个Tab完整Web工作台；LaTeX报告导出；完整审计链导出；23帧设计稿验收'],
            ['代码/结果与流程可复现性', '10', 'Mock LLM模式；一键启动脚本；14项E2E验收；7批回归测试；jsonl审计链'],
        ])

    # ===== CHAPTER 10: Summary =====
    add_heading1(doc, '十、总结与展望')
    add_body(doc, '"联邦智研"（AISci）项目基于国产大模型Qwen，构建了一套完整的多智能体科研自动化系统。系统以科学逻辑链为核心，通过7阶段Pipeline覆盖"研究问题→文献挖掘→知识缺口→假设生成→假设评审→迭代实验→报告生成"的全流程，并以证据链迭代推理引擎、布尔质量门禁系统、统一反馈中心和反事实预演机制四大创新为技术支撑，有效解决了LLM在科研场景中的幻觉引用、质量控制和闭环迭代三大核心挑战。')
    add_body(doc, '项目的技术深度体现在：6个独立Agent的标准化设计、约70个可复用Skill模块的领域覆盖、11种阶段特定布尔Gate的精密度量、以及科学自迭代的闭环编排能力。项目的应用价值体现在：面向Science 125科学问题的实证验证、联邦学习领域的扩展支持、以及完整的前端工作台和审计链导出能力，确保了系统的可用性和可复现性。')
    add_body(doc, '未来展望：①引入更多科学领域的数据集和领域专家知识，扩展系统的跨学科覆盖范围；②将L0定性反事实预演升级为L1定量仿真，提升实验方案评估的精确度；③与真实实验室仪器环境对接，实现从虚拟仿真到物理实验的闭环；④探索多Agent间基于强化学习的协作策略优化，实现更高效的任务分配和资源调度。')

    # ===== APPENDIX =====
    doc.add_page_break()
    add_heading1(doc, '附录：系统功能示意图')
    add_body(doc, '图1-1：AISci系统功能示意图（待插入designs/aisci-functional-schematic.html导出图片）')
    add_body(doc, '图2-1：系统整体架构图（待插入架构图）')
    add_body(doc, '图3-1：七阶段Pipeline流程图')
    add_body(doc, '图4-1：证据链迭代推理引擎流程图')
    add_body(doc, '图5-1：前端12 Tab工作台截图（23帧设计稿）')

    # ===== Save =====
    output_dir = 'D:/Workplace/AISci/output'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, '联邦智研_大创申报书.docx')
    doc.save(output_path)
    print(f'申报书已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    main()
