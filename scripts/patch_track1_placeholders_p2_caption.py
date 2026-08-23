# -*- coding: utf-8 -*-
"""删除已作答的题干占位句，并为 P2 总体思路图补图注。"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

sys.path.insert(0, str(Path(__file__).resolve().parent))
from restore_track1_underlines import _make_run, _wipe_runs  # noqa: E402
from retune_track1_1a_theme import set_fill  # noqa: E402

TPL = Path(
    r"d:/Workplace/AISci/output/提交/模板/赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx"
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

P2_CAPTION = (
    "图 P2　总体思路。实线为七阶段主链路（问题理解→文献挖掘→知识缺口→假设生成→假设评审"
    "→迭代实验→报告生成）；虚线为三大创新注入与反馈：证据链与 Fact 白名单、对齐与预检门禁、"
    "「大家长」审核各阶段输入输出并启停证据链迭代。"
)

P19_BOUNDARY = (
    "候选假设和研究计划仍需研究者审查，不能等同于已经获得科学发现或完成真实验证："
    "受大模型用量和本地算力限制，现在交出来的是可以继续往下验证的研究方案（更接近开题材料），"
    "不能当作已经得到科学发现，也不能当作真实验证已经做完。"
)


def insert_paragraph_after(paragraph, text: str, underlined: bool = True):
    new_p = etree.Element("{%s}p" % W_NS)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, paragraph._parent)
    para._p.append(_make_run(text, underline=underlined))
    return para


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def main() -> None:
    doc = Document(str(TPL))

    # 1) 删除纯题干占位句（答案已在下一段）
    to_delete = []
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("[请选择") or t.startswith("[请说明"):
            to_delete.append(i)
            print("will delete placeholder", i, t[:60])
    for i in sorted(to_delete, reverse=True):
        delete_paragraph(doc.paragraphs[i])

    # 2) P19：合并「请明确…」题干与「答：…」
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("请明确候选假设和研究计划仍需研究者审查"):
            delete_paragraph(p)
            print("deleted P19 prompt", i)
            break
    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        if t.startswith("答：受大模型用量"):
            set_fill(p, P19_BOUNDARY)
            print("rewrote P19 answer", i)
            break

    # 3) P2 图注：插在总体思路图后（有 drawing、下一段不是「图 P2」）
    inserted = False
    for i, p in enumerate(doc.paragraphs):
        if "w:drawing" not in p._p.xml:
            continue
        prev = (doc.paragraphs[i - 1].text or "") if i else ""
        if "总体思路" not in prev:
            continue
        nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
        nxt_text = (nxt.text or "") if nxt is not None else ""
        if nxt_text.startswith("图 P2"):
            _wipe_runs(nxt._p)
            nxt._p.append(_make_run(P2_CAPTION, underline=True))
            print("updated existing P2 caption", i + 1)
        else:
            insert_paragraph_after(p, P2_CAPTION, underlined=True)
            print("inserted P2 caption after", i)
        inserted = True
        break
    if not inserted:
        raise RuntimeError("P2 figure paragraph not found")

    # 4) 顺带统一 P7 图注空格（「图 P7一次」→「图 P7　一次」）
    for p in doc.paragraphs:
        t = p.text or ""
        if t.startswith("图 P7一次"):
            set_fill(p, t.replace("图 P7一次", "图 P7　一次", 1))
            print("normalized P7 caption spacing")
            break

    try:
        doc.save(str(TPL))
        out = TPL
    except PermissionError:
        out = TPL.with_name(TPL.stem + "._caption_patched.docx")
        doc.save(str(out))
        print("LOCKED, wrote", out.name)
        return

    print("saved", out.name)

    doc2 = Document(str(out))
    left = [p.text for p in doc2.paragraphs if (p.text or "").strip().startswith("[请")]
    print("remaining [请 placeholders]", len(left), left)
    for i, p in enumerate(doc2.paragraphs):
        t = p.text or ""
        if t.startswith("图 P2"):
            print("P2 caption:", t[:100])
        if "w:drawing" in p._p.xml:
            prev = (doc2.paragraphs[i - 1].text or "")[:20] if i else ""
            nxt = (doc2.paragraphs[i + 1].text or "")[:40] if i + 1 < len(doc2.paragraphs) else ""
            if "总体思路" in prev:
                print("P2 pic next=", nxt)


if __name__ == "__main__":
    main()
