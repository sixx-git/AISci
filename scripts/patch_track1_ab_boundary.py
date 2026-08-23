# -*- coding: utf-8 -*-
"""补强 P12 反馈表本案路径，并按「只交 1A；迭代实验用实验闭环设计服务 1A」改 P13/P17。"""
from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from retune_track1_1a_theme import set_cell, set_fill  # noqa: E402

TPL = Path(
    r"d:/Workplace/AISci/output/提交/模板/赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx"
)

# P12 反馈表新增本案主路径
P12_NEW_ROW = [
    "计划不可执行（证据等级仍可为 high）",
    "入选假设的验证写成万例前瞻、纠纷终点等，当前做不了；与「证据弱」不是同一类问题。",
    "大家长终止当轮，交 HITL；不改写已审核的第一轮输出。后续在实验页按『任务规划—运行—数据—反馈』闭环调整（服务于 1A 的验证，不另交实验赛道）。",
    "人工写入收窄约束后 fork 第二轮证据链；旧实验指标不继承给新假设。",
]

# P13：125 题测试方法 —— 改「为填写模板手动重跑」口径
P13_125 = (
    "125 题测试用来检验方法链能否跨学科复用，每题单次跑完七阶段，并进行一次非全量的迭代实验预检，"
    "统一检查四项产出：可处理缺口、带反对证据的候选假设、筛选与入选理由、可检验研究计划；"
    "并保留审计链以便对照第一轮。"
    "候选假设生成后经三级评审：LLM 结构化评审按科学价值、新颖性、可检验性、数据可得性、成本风险五维打分；"
    "多评审者集成评审聚合导师模拟、新颖性 Skill 与证据规则三类评审者加权得出综合分（≥6.5 入选）；"
    "评审分歧或低分时进入人工评审，选择某个假设或指定某阶段重跑。"
    "任务打断后，用户按提示完成操作再继续。"
    "工程上 120/125 一次跑通，5 题因接口外部失败重试后完成，题目本身没有被省略；"
    "发生证据不足或阶段失败时，由 Agent 提示后自动或人工做某阶段重启。"
    "质量核验：引用核验通过（可疑引用 0）；假设能回溯高证据事实约 56.2%；"
    "领域对齐约 79.1%；报告逻辑审查均分约 6.08 分（满分 10 分）。"
    "代表案例另选 sjtu_q_087 展示 1A 要求的自迭代："
    "多数题目单次跑完即可，证据不足时做某阶段重启，不必开完整第二轮证据链；"
    "本案在第一轮已冻结（run_id=c6e7fbfb-…）之后，按人工意见另开第二轮分支（run_id=f2c80a5f-…），"
    "用于版本比较与反馈修正。"
    "「手动」指人工触发并确认从文献挖掘起重跑、保留两轮对照，不是事后改写第一轮冻结件。"
)

# P2 实际完成：补一句 A 主交 / 实验闭环服务 A（避免与「影响力 1B 延伸」混淆）
P2_WORK = (
    "已用三大创新作为实现手段，在七阶段中把 1A 要证明的四步跑出来。"
    "创新一：文献事实进入 Fact 白名单，绑定 chunk_id，并做反证检索与立场标注。"
    "创新二：通过知识库对齐、集成评审、可证伪性过滤、可执行性门禁与小样本预检，留下筛选理由并过滤空泛计划。"
    "创新三：「大家长」审核各阶段输入输出，证据不足则启动证据链迭代（补文献/重跑），"
    "收敛或需人判断则终止并交 HITL 复核；下一轮只追加约束，不改写上一轮已审核的输出。"
    "问题理解与知识缺口识别对应要证明的第一步，其证据覆盖度来自创新一产出的事实白名单。"
    "本作品按方向 1A（科学假设生成与研究计划设计）提交；"
    "迭代实验页按『任务规划—实验运行—数据分析—反馈迭代』闭环实现，"
    "用于兑现 1A 命题中的验证与反馈修正，不作为单独方向证明。"
    "125 题用于检验这套手段能否跨学科跑通；科学影响力预测属延伸能力，不作为 1A 主证明。"
)

# P17 说明行
P17_NOTE = (
    "说明：本作品只提交方向 1A。"
    "证据链迭代（寻找更好假设、版本比较）与迭代实验（对已入选假设做任务规划—运行—数据—反馈，"
    "记录于 09_iterative_experiments.json）是两条独立线："
    "前者兑现 1A 的「假设从何而来 / 依据什么 / 如何自迭代」；"
    "后者采用命题中实验闭环的设计，兑现 1A 的「可能如何验证 / 反馈修正」，不另交实验赛道。"
    "因此第二轮证据链收窄假设后，不自动产生新的完整研究计划，也不把第一轮实验指标贴到新假设上；"
    "上表只对照证据链带来的假设变化，迭代实验明细见下表。"
)

P17_IMPROVE = (
    "第二轮实际改善了什么：在 1A 证据链一侧，入选假设的验证落到公开影像上的不确定性门控。"
    "新 H-01 可检验性原文：「在公开医学影像数据集上训练带MC-Dropout的CNN模型，"
    "设定不同不确定性阈值模拟人机协同决策流，……优先使用CheXpert或ISIC公开库。」"
    "反对证据与「最终责任仍由执业医师承担」写得更明确；评审 overall 7.04→7.29。"
    "改善的是假设可检验性与版本可对照，不是把第一轮改成从未选错；"
    "1A 的「四步」由第一轮完整中间产物与第二轮对照合起来证明，不是每轮各再跑一遍四步。"
)

P17_COST = (
    "哪些方面没有改善或出现了新的代价：真实前瞻仍未做；"
    "第二轮未在实验页按新 H-01 重开『任务规划—运行—反馈』闭环，"
    "故第一轮 RSNA 烟雾测试（250 行，silhouette 0.3219 等）不能当作新 H-01 的验证——"
    "这是两条迭代线未在第二轮汇合的边界，不是否认实验闭环能力；"
    "仍有一条依赖约 200 名医师模拟交互的候选未入选。"
)

P17_STOP = (
    "团队为什么在此时停止，或为什么仍需继续迭代："
    "停在第二轮评审后的 HITL，是为了并排对照两轮假设与第一轮实验，证明没有覆盖第一轮。"
    "若继续，应在实验页按新 H-01 启动任务规划与仿真验证（实验闭环服务于 1A），"
    "而不是把旧指标改贴到新假设上，也不是再开一轮只为「看起来又跑完四步」。"
)


def add_table_row(table, cells_text: list[str]) -> None:
    tbl = table._tbl
    last_tr = table.rows[-1]._tr
    new_tr = deepcopy(last_tr)
    tbl.append(new_tr)
    new_row = table.rows[-1]
    for cell, text in zip(new_row.cells, cells_text):
        set_cell(cell, text)


def patch_para_startswith(doc: Document, prefixes: str | tuple[str, ...], new_text: str) -> int:
    if isinstance(prefixes, str):
        prefixes = (prefixes,)
    for i, p in enumerate(doc.paragraphs):
        text = p.text or ""
        if any(text.startswith(prefix) for prefix in prefixes):
            set_fill(p, new_text)
            return i
    raise RuntimeError(f"paragraph not found: {prefixes[0][:40]}")


def main() -> None:
    doc = Document(str(TPL))

    t12 = doc.tables[10]
    already = any("计划不可执行" in (r.cells[0].text or "") for r in t12.rows)
    if not already:
        add_table_row(t12, P12_NEW_ROW)
        print("added P12 feedback row")
    else:
        for r in t12.rows:
            if "计划不可执行" in (r.cells[0].text or ""):
                for ci, val in enumerate(P12_NEW_ROW):
                    set_cell(r.cells[ci], val)
                print("updated existing P12 feedback row")
                break

    print("P2", patch_para_startswith(doc, "已用三大创新作为实现手段", P2_WORK))
    print("P13", patch_para_startswith(doc, "125 题测试用来检验方法链", P13_125))
    print(
        "P17note",
        patch_para_startswith(
            doc,
            ("说明：本作品只提交方向 1A", "说明：迭代实验"),
            P17_NOTE,
        ),
    )
    print("P17imp", patch_para_startswith(doc, "第二轮实际改善了什么：", P17_IMPROVE))
    print("P17cost", patch_para_startswith(doc, "哪些方面没有改善或出现了新的代价：", P17_COST))
    print("P17stop", patch_para_startswith(doc, "团队为什么在此时停止", P17_STOP))

    out = TPL
    try:
        doc.save(str(out))
    except PermissionError:
        out = TPL.with_name(TPL.stem + "._ab_boundary_patched.docx")
        doc.save(str(out))
        print("LOCKED, wrote", out.name)
        return

    print("saved", out.name)

    doc2 = Document(str(out))
    print("T10 rows", len(doc2.tables[10].rows))
    print("last row", doc2.tables[10].rows[-1].cells[0].text[:40])
    checks = [
        ("P17 note OK", "只提交方向 1A"),
        ("P13 OK", "「手动」指人工触发"),
        ("P2 OK", "不作为单独方向证明"),
        ("P17 four-step OK", "不是每轮各再跑一遍四步"),
    ]
    blob = "\n".join(p.text or "" for p in doc2.paragraphs)
    for label, needle in checks:
        print(label if needle in blob else f"MISSING {label}")


if __name__ == "__main__":
    main()
