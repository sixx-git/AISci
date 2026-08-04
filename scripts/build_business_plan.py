# -*- coding: utf-8 -*-
"""Generate 联邦智研 AISci 商业版申报书 (互联网+/创新创业方向) with python-docx.
Follows competition-doc-standards: A4, 2.54cm margins, 宋体, Heading 1/2/3,
body 1.5 line spacing + 2-char first-line indent, Table Grid + #CFCDCD header.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r"D:\Workplace\AISci\output\项目计划书-联邦智研AISci-商业版.docx"

FONT = "宋体"
HEADER_FILL = "CFCDCD"

align_map = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def set_run_font(run, font_name=FONT, size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    rfonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell(cell, text, size=10.5, bold=False, align='center', fill=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.line_spacing = 1.2
    for i, line in enumerate(str(text).split('\n')):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, FONT, size, bold)
    if fill:
        set_cell_shading(cell, fill)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, FONT, [0, 18, 15, 13][level], bold=True)
    return p


def add_body(doc, text, indent=True, size=12, bold=False, align='justify'):
    p = doc.add_paragraph()
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Cm(0.85)
    pf.space_after = Pt(6)
    for i, line in enumerate(str(text).split('\n')):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, FONT, size, bold)
    return p


def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.85)
    pf.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, FONT, size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, size=10.5, bold=True, align='center', fill=HEADER_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val, size=10.5, bold=False, align='left')
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2); run._r.append(fld3)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(instr); run._r.append(f2)
    set_run_font(run, FONT, 9)


# ---------------------------------------------------------------------------
MEDIA_DIR = r"D:\Workplace\AISci\output\_v3_media"


def add_figure(doc, img_file, caption):
    """Insert a centered figure with caption below."""
    path = os.path.join(MEDIA_DIR, img_file)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(15))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(cap.add_run(caption), FONT, 10)
    cap.paragraph_format.space_after = Pt(12)


# ---------------------------------------------------------------------------
doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(12)
nrpr = normal.element.get_or_add_rPr()
nf = nrpr.find(qn('w:rFonts'))
if nf is None:
    nf = OxmlElement('w:rFonts'); nrpr.append(nf)
nf.set(qn('w:eastAsia'), FONT)

for sname, sz in [('Heading 1', 18), ('Heading 2', 15), ('Heading 3', 13)]:
    st = doc.styles[sname]
    st.font.name = FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), FONT)

# Page setup A4, 2.54cm margins
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(2.54)
sec.right_margin = Cm(2.54)
add_page_number_footer(sec)

# ---------------- Cover ----------------
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('联邦智研 AISci'); set_run_font(r, FONT, 26, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('面向科研创新全流程的 AI 科研智能体平台'); set_run_font(r, FONT, 15, bold=False)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('大学生创新创业大赛 · 商业计划书'); set_run_font(r, FONT, 14, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('（互联网+ / 高教主赛道 · 人工智能与大数据方向）'); set_run_font(r, FONT, 11)
for _ in range(4):
    doc.add_paragraph()
for line in ['项目名称：联邦智研 AISci（科研自动化智能体系统）',
             '参赛赛道：高教主赛道 / 人工智能与大数据',
             '团队名称：联邦智研团队',
             '编制日期：2026 年']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(line), FONT, 12)

doc.add_page_break()

# ---------------- TOC ----------------
add_heading(doc, '目  录', 1)
add_toc(doc)
doc.add_page_break()

# ---------------- 第一章 项目概述 ----------------
add_heading(doc, '第一章 项目概述', 1)
add_heading(doc, '1.1 项目背景与时代机遇', 2)
add_body(doc, '随着人工智能加速进入科学研究领域，AI for Science 已成为全球科技竞争的战略制高点。然而，科研人员正面临"信息爆炸"与"认知瓶颈"的双重挑战：全球每年发表学术论文超过 300 万篇，单点研究者难以全面追踪前沿；大型语言模型（LLM）虽具备强大文本理解与生成能力，却普遍存在"幻觉引用""缺乏证据可追溯性""缺少系统化质量控制"三大痛点，难以直接支撑严谨的科研活动。')
add_body(doc, '在此背景下，本项目"联邦智研 AISci"基于国产大模型 Qwen，打造面向高校、科研机构与企业研发团队的科研智能体系统，将"问题—证据—假设—验证—优化"的科研逻辑链编码为可自动执行的多智能体 Pipeline，使 AI 从"问答工具"升级为"可信的科研协作者"。')

add_heading(doc, '1.2 项目定位与愿景', 2)
add_body(doc, '我们刻意区别于市场上大量"AI 科研助手"的定位。AISci 的准确定位是：')
add_bullet(doc, '科研智能体基础设施（Research Agent Infrastructure）——为科研全流程提供可编排、可审计、可复现的智能能力底座；')
add_bullet(doc, '面向高校和研发机构的 AI 科研生产力平台——把分散的文献、假设、实验与报告环节连成闭环。')
add_body(doc, '愿景：让每一个科研团队，无论规模大小，都拥有可信、可控、可解释的 AI 科研协作者，把研究者从重复性劳动中解放出来，聚焦真正具有创造力的科学判断。')

add_heading(doc, '1.3 核心能力一句话', 2)
add_body(doc, '通用 AI（如 ChatGPT）的能力是"问题 → 回答"；AISci 提供的是"问题 → 证据 → 假设 → 验证 → 优化"的科研闭环。这一差异，正是本项目商业价值的根本来源。')

add_heading(doc, '1.4 项目亮点速览', 2)
add_bullet(doc, '系统完整：已形成包含前端、后端、API、测试与可复现性保障的近产品原型，而非概念演示；')
add_bullet(doc, '国产适配：基于 Qwen 构建，可无缝对接国产 AI 基础设施与信创环境；')
add_bullet(doc, '可信机制：以 Evidence Chain、Boolean Gate、HITL 三位一体解决 AI 幻觉与质量不可控问题；')
add_bullet(doc, '全流程闭环：覆盖从科学问题到可验证假设、再到实验迭代与报告生成的完整链路。')

# ---------------- 第二章 市场痛点 ----------------
add_heading(doc, '第二章 市场痛点与需求分析', 1)
add_heading(doc, '2.1 科研人员效率瓶颈', 2)
add_body(doc, '高校与科研机构科研人员大量时间消耗在文献阅读、资料整理与格式规范上；青年教师普遍缺乏系统的科研方向发现能力；跨学科研究因知识壁垒高、检索成本高而难以开展。')
add_heading(doc, '2.2 创新发现困难', 2)
add_body(doc, '大量科研机会隐藏于跨领域知识的交叉区域，但人工难以系统扫描学科边界；研究者容易陷入"已有工作的重复"，错失真正的创新空白点。')
add_heading(doc, '2.3 企业研发成本高', 2)
add_body(doc, '企业研发部门面临专利检索成本高、技术路线分析困难、新产品研发周期长等压力；尤其在医药、新材料、AI 等强研发行业，方向判断失误的代价极为高昂。')
add_heading(doc, '2.4 科研机构管理低效', 2)
add_body(doc, '科研院所与高校科研管理部门存在项目评审效率低、科研成果评价困难、学科态势感知滞后等问题，缺乏智能化的科研治理工具。')
add_heading(doc, '2.5 目标用户画像与需求', 2)
add_table(doc,
    ['用户群体', '核心痛点', '核心需求', '付费意愿'],
    [
        ['高校科研人员\n（师生/团队）', '文献量大、方向难寻、跨学科困难', 'AI 辅助文献分析、研究方向发现、项目申请辅助', '中（个人/课题组订阅）'],
        ['企业研发部门\n（药企/新材料/AI）', '专利检索贵、技术路线不清、周期长', '技术趋势分析、专利分析、创新方案生成', '高（年费/项目制）'],
        ['科研机构\n（高校/院所管理）', '评审低效、成果评价难、态势感知弱', '智能科研管理、项目评估、学科分析', '高（私有化部署）'],
    ],
    col_widths=[3.2, 5.0, 5.5, 2.8])

# ---------------- 第三章 产品体系 ----------------
add_heading(doc, '第三章 产品体系', 1)
add_body(doc, 'AISci 围绕三类核心用户，构建三个相互协同的产品线，形成从个人科研到机构治理的完整覆盖。')
add_heading(doc, '3.1 AISci Research（面向高校科研人员）', 2)
add_bullet(doc, 'AI 文献助手：语义检索、证据抽取、综述生成；')
add_bullet(doc, '研究方向发现：基于知识缺口自动推荐潜在研究空白；')
add_bullet(doc, '科研计划生成：一键产出含假设、方法与实验路线的研究方案。')
add_heading(doc, '3.2 AISci Enterprise（面向企业研发）', 2)
add_bullet(doc, '技术趋势分析：追踪行业技术演进与竞争态势；')
add_bullet(doc, '专利分析：查新、规避设计与自由实施（FTO）辅助；')
add_bullet(doc, '创新方案生成：针对具体研发目标产出可验证的技术方案。')
add_heading(doc, '3.3 AISci Institution（面向高校与科研机构）', 2)
add_bullet(doc, '科研管理：项目全生命周期跟踪与里程碑预警；')
add_bullet(doc, '项目评估：基于证据链的客观化评审支持；')
add_bullet(doc, '学科分析：学科竞争力与交叉机会透视。')

# ---------------- 第四章 技术与核心竞争力 ----------------
add_heading(doc, '第四章 产品技术与核心竞争力', 1)
add_heading(doc, '4.1 系统整体架构', 2)
add_body(doc, '系统采用前后端分离的三层架构，遵循"Agent—Skill—Infrastructure"三层分离原则：后端基于 Python FastAPI 构建，提供 18 个 API 路由模块；前端基于 React 18 + Vite 5 + TailwindCSS 构建；数据层使用 SQLite 配合 zvec 向量检索，支持文献全文的语义搜索与相似度匹配。该架构具备良好的可部署性与横向扩展能力，可支撑 SaaS、私有化等多种交付形态。')
add_figure(doc, 'image2.png', '图4-1 AISci 系统整体架构图（三层设计）')
add_heading(doc, '4.2 七阶段科研 Pipeline', 2)
add_body(doc, '系统核心是七阶段顺序 Pipeline，每个阶段由独立 Agent 负责，通过标准化 JSON Schema 保证阶段间数据传递的类型安全：')
add_body(doc, '问题理解 → 文献挖掘 → 知识缺口发现 → 假设生成 → 假设评审 → 迭代实验 → 报告生成。\n每个阶段执行后自动记录完整审计数据（输入、输出、Prompt 原文、模型参数、Token 用量、耗时及每次 API 调用记录），支持 jsonl 导出，确保科研过程完全可复现。', align='left')
add_figure(doc, 'image3.png', '图4-2 AISci 七阶段 Pipeline 流程图')
add_heading(doc, '4.3 六大核心智能体', 2)
add_table(doc,
    ['智能体', '职责'],
    [
        ['问题理解 Agent', '将模糊科学问题转化为结构化描述（主要矛盾、研究对象三维拆解、研究目的）'],
        ['文献挖掘 Agent', '向量召回 + LLM 证据抽取，每条事实绑定来源 chunk_id，支持 arXiv 与 PDF 解析'],
        ['知识缺口 Agent', '分析已有研究覆盖，识别未被充分研究的空白区域与可突破机会'],
        ['假设生成 Agent', '基于证据生成候选假设并标注 supporting_fact_ids，Pairwise 锦标赛排序'],
        ['假设评审 Agent', '多维评估新颖性/可行性/反事实预演，Ensemble Gate 判定通过'],
        ['迭代实验 Agent', '绑定数据→设计脚本→沙箱执行→结果分析→重设计，形成闭环'],
    ],
    col_widths=[3.5, 13.0])
add_heading(doc, '4.4 技能工具层（Skill Layer）', 2)
add_body(doc, 'Skill 层是 Agent 的可复用工具库，按科研子领域组织为 8 大类、约 70 个工具模块，包括文献检索类、PDF 解析类、证据链类、假设生成类、实验设计类、质量门禁类、反馈类与报告类。Agent 负责流程编排与 LLM 调用，Skill 提供具体研究工具能力，二者解耦使系统具备强扩展性与可维护性。')
add_heading(doc, '4.5 四大核心创新机制与商业价值映射', 2)
add_body(doc, '项目在机制层面的创新，最终都指向明确的客户价值。下表将技术语言翻译为商业语言，便于评委与用户直接理解：')
add_table(doc,
    ['技术创新', '解决的问题', '为客户带来的商业价值'],
    [
        ['Evidence Chain\n（证据链迭代推理）', 'LLM 幻觉引用、证据不可追溯', '降低 AI 生成错误引用风险，提升产出可信度'],
        ['Boolean Gate\n（布尔质量门禁）', '连续评分缺乏决策支持', '以硬门禁替代模糊评分，提升科研结果可信度'],
        ['Feedback Hub + HITL\n（统一反馈中心+人机协同）', '阶段间反馈割裂、失控', '降低人工审核成本，平衡自动化与可控性'],
        ['Counterfactual\n（反事实预演）', '实验方向盲目、资源浪费', '提升实验方案成功率，保护研发资源'],
    ],
    col_widths=[4.2, 4.8, 7.5])
add_figure(doc, 'image4.png', '图4-3 证据链迭代推理引擎流程图（核心创新一）')
add_heading(doc, '4.6 科学自迭代闭环与可复现性', 2)
add_body(doc, '当证据强度不足或评审未通过时，系统自动触发"证据弱→补文献→重建证据链→重跑假设树→再评审"的自迭代闭环，并记录完整因果链摘要。全流程审计数据可导出，支撑科研结果的完全可复现——这正是企业研发与科研机构最为看重的"可追溯、可审计"能力。')

# ---------------- 第五章 竞品分析 ----------------
add_heading(doc, '第五章 竞品分析', 1)
add_heading(doc, '5.1 国内外竞品对比', 2)
add_table(doc,
    ['产品', '定位', '优势', '不足'],
    [
        ['ChatGPT', '通用 AI 助手', '语言能力强、用户广', '无科研闭环、易幻觉'],
        ['Elicit', '文献分析', '文献检索强', '无假设生成、无实验'],
        ['NotebookLM', '知识库问答', '文档理解好', '无实验规划、无闭环'],
        ['AI Scientist', '自动科研', '可自动生成论文', '可控性不足、易幻觉假设'],
        ['AISci（本项目）', '科研智能体', '全流程闭环、可信可控', '需市场验证、品牌待建'],
    ],
    col_widths=[2.8, 3.0, 4.2, 6.5])
add_heading(doc, '5.2 核心能力对比', 2)
add_table(doc,
    ['能力', 'ChatGPT', 'Elicit', 'NotebookLM', 'AI Scientist', 'AISci'],
    [
        ['文献分析', '√', '√', '√', '△', '√'],
        ['假设生成', '△', '×', '×', '√', '√'],
        ['证据链', '×', '△', '△', '×', '√'],
        ['科研闭环', '×', '×', '×', '△', '√'],
        ['实验规划', '×', '×', '×', '√', '√'],
    ],
    col_widths=[3.0, 2.7, 2.7, 3.0, 3.0, 2.6])
add_body(doc, '结论：竞品多在局部能力取得进展，AISci 是首个将"自动假设生成 + 证据可追溯 + 人在回路 + 闭环自迭代"统一于一体的科研自动化系统。我们的差异化优势不是"模型更大"，而是"科研流程更完整、结果更可信"。', align='left')

# ---------------- 第六章 商业模式 ----------------
add_heading(doc, '第六章 商业模式', 1)
add_heading(doc, '6.1 高校版 SaaS', 2)
add_body(doc, '客户：高校实验室与科研团队。收费：1,999 ~ 9,999 元 / 年 / 账号。功能：文献分析、研究方向发现、项目申请辅助。以低成本订阅降低高校用户决策门槛，快速积累口碑与标杆案例。')
add_heading(doc, '6.2 企业研发版', 2)
add_body(doc, '客户：医药企业、新材料企业、AI 企业等研发密集型企业。收费：10 万 ~ 100 万元 / 年。功能：技术趋势分析、专利分析、创新方案生成。以高客单价支撑规模化营收。')
add_heading(doc, '6.3 私有化部署', 2)
add_body(doc, '客户：高校、科研院所等数据敏感机构。收费：项目制（含部署、定制与运维）。满足信创与数据不出域的合规要求，建立长期深度合作。')
add_heading(doc, '6.4 盈利模式与定价策略', 2)
add_table(doc,
    ['产品', '目标客户', '收费模式', '核心功能'],
    [
        ['AISci Research', '高校师生/团队', 'SaaS 订阅 1999~9999 元/年/账号', '文献助手、方向发现、计划生成'],
        ['AISci Enterprise', '企业研发部门', '年费 10~100 万元', '技术趋势、专利分析、方案生成'],
        ['AISci Institution', '高校/科研院所', '私有化项目制', '科研管理、项目评估、学科分析'],
    ],
    col_widths=[3.2, 3.3, 4.5, 5.5])

# ---------------- 第七章 技术实现与工程化 ----------------
add_heading(doc, '第七章 技术实现与工程化保障', 1)
add_heading(doc, '7.1 前后端与 API', 2)
add_body(doc, '后端 FastAPI 提供 18 个 API 路由模块，前端通过 PipelineProgress 组件实时展示进度、RunLogDetail 查看每步详情，实现科研过程全透明化与可追溯化；HITL 机制支持用户在任意阶段查看、编辑、重跑或与阶段输出交互式对话。')
add_heading(doc, '7.2 测试覆盖与可复现性', 2)
add_body(doc, '系统配套 pytest 测试体系与完整环境配置文档，全链路审计数据支持 jsonl 导出，确保第三方可在相同条件下复现实验结果——这是面向科研机构与企业客户的关键信任基础。')
add_heading(doc, '7.3 代表性测试案例与验证结果', 2)
add_bullet(doc, '案例一（Science 125 问题）：选取 10 个跨学科典型问题运行完整 Pipeline，假设新颖性评分 ≥ 6.0、证据链至少 1 轮迭代、References 无虚构引用；')
add_bullet(doc, '案例二（评分表系统）：6 组测试均生成完整评分表，总评分 88–118 分，维度 41–56 项，稳定性与区分度良好；')
add_bullet(doc, '案例三（联邦学习仿真）：支持 local_pack / Flower / FedML 三种后端，验证实验脚本可执行性。')

# ---------------- 第八章 竞争优势与发展规划 ----------------
add_heading(doc, '第八章 竞争优势与发展规划', 1)
add_heading(doc, '8.1 核心竞争优势', 2)
add_table(doc,
    ['能力', 'ChatGPT', 'AISci'],
    [
        ['文献分析', '√', '√'],
        ['假设生成', '△', '√'],
        ['证据链', '×', '√'],
        ['科研闭环', '×', '√'],
        ['实验规划', '×', '√'],
    ],
    col_widths=[4.0, 4.0, 4.0])
add_body(doc, 'AISci 的护城河不是单一模型能力，而是"科研流程完整性 + 可信机制 + 国产适配"的组合。通用大模型可作为底层能力被集成，但科研闭环与质量治理是难以被简单复制的系统性壁垒。', align='left')
add_heading(doc, '8.2 三阶段发展规划', 2)
add_table(doc,
    ['阶段', '时间', '重点', '目标'],
    [
        ['第一阶段', '0–1 年', '高校实验室试点、产品打磨', '100 个种子用户，沉淀标杆案例'],
        ['第二阶段', '1–3 年', '切入企业研发市场', '20 家企业客户，建立年费营收'],
        ['第三阶段', '3–5 年', '成为 AI 科研基础设施平台', '覆盖科研全链条，形成生态'],
    ],
    col_widths=[2.8, 2.8, 5.2, 5.7])

# ---------------- 第九章 融资规划 ----------------
add_heading(doc, '第九章 融资规划', 1)
add_heading(doc, '9.1 融资需求与资金用途', 2)
add_body(doc, '本轮融资主要用于模型优化、数据建设、市场推广与产品商业化，加速从"优秀科研系统"向"具有产业化潜力的 AI 创业项目"跃迁。')
add_heading(doc, '9.2 资金分配', 2)
add_table(doc,
    ['用途', '占比', '说明'],
    [
        ['模型优化', '30%', 'Qwen 微调、证据链与门禁算法迭代'],
        ['数据建设', '20%', '文献库、基准数据集、行业语料'],
        ['市场推广', '30%', '高校地推、企业 BD、品牌建设'],
        ['产品商业化', '20%', 'SaaS 平台、私有化交付能力'],
    ],
    col_widths=[3.5, 2.5, 10.5])

# ---------------- 第十章 团队与风险 ----------------
add_heading(doc, '第十章 团队与风险', 1)
add_heading(doc, '10.1 团队优势', 2)
add_body(doc, '团队具备 AI 算法、科研领域知识与工程化落地能力的复合背景，并在导师指导下完成系统从 0 到 1 的构建，已具备完整前后端与可复现性保障。')
add_heading(doc, '10.2 风险识别与对策', 2)
add_table(doc,
    ['风险类型', '具体风险', '对策'],
    [
        ['技术风险', '模型幻觉、长链路稳定性', 'Evidence Chain + Boolean Gate 硬约束，持续迭代'],
        ['市场风险', '用户付费习惯未建立', '高校低价 SaaS 引流，企业高客单转化'],
        ['竞争风险', '大厂下场做科研 AI', '聚焦科研闭环壁垒与国产适配差异化'],
        ['合规风险', '数据隐私与版权', '私有化部署、数据不出域、文献合规使用'],
    ],
    col_widths=[3.0, 5.5, 7.0])

# ---------------- 附录：可视化支撑材料 ----------------
doc.add_page_break()
add_heading(doc, '附录 可视化支撑材料', 1)
add_body(doc, '以下为系统核心可视化示意图，展示 AISci 的完整功能架构、前端交互界面与核心引擎设计。', indent=False, size=11)
add_figure(doc, 'image1.png', '图A-1 AISci 系统功能架构（前端交互层 + 核心引擎层 + 数据与输出层）')

doc.save(OUT)
print('SAVED:', OUT)
