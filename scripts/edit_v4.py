# -*- coding: utf-8 -*-
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PATH = "D:/Workplace/AISci/output/联邦智研_大创申报书_v4.docx"
doc = Document(PATH)


def set_run_font(run, cn='宋体', en='宋体', size=None, bold=None):
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), cn)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)


def new_para(ref, text=None, style=None, before=False, align=None,
             first_line=None, left=None, ls=1.5, sa=Pt(6), font='宋体', bold=None):
    p = doc.add_paragraph()
    if style is not None:
        p.style = style
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if first_line is not None:
        pf.first_line_indent = first_line
    if left is not None:
        pf.left_indent = left
    pf.line_spacing = ls
    pf.space_after = sa
    if text is not None:
        r = p.add_run(text)
        if font:
            set_run_font(r, cn=font, en=font, bold=bold)
    if before:
        ref._p.addprevious(p._p)
    else:
        ref._p.addnext(p._p)
    return p


def block_before(ref, items):
    anchor = ref
    for text, kw in reversed(items):
        p = new_para(anchor, text=text, before=True, **kw)
        anchor = p
    return anchor


# ---- collect references (object refs survive edits) ----
ch1 = doc.paragraphs[11]            # 一、研究问题与解决方法
p21_end = doc.paragraphs[28]        # end of 2.1 body
p_pipeline_arrow = doc.paragraphs[32]  # arrow listing in 2.2
p31_intro = doc.paragraphs[42]      # intro of 3.1
p41_end = doc.paragraphs[72]        # end of 4.1
p12_end = doc.paragraphs[21]        # end of 1.2
p42_end = doc.paragraphs[76]
p43_end = doc.paragraphs[80]
p44_end = doc.paragraphs[83]
p71_end = doc.paragraphs[121]       # end of 7.1
ch8_head = doc.paragraphs[132]      # 八、科学影响力预测系统
p82_end = doc.paragraphs[141]       # end of 8.2
detail_paras = list(doc.paragraphs[142:174])   # 8.3 - 8.8 (32 paras)
appx_img4 = doc.paragraphs[221]     # 图4-1 image in appendix
para19 = doc.paragraphs[19]
para85 = doc.paragraphs[85]
H1 = ch1.style                      # chapter heading style
H2 = doc.paragraphs[26].style       # sub-heading style

# ============ ITEM 1 + 4 + 5 : 摘要 (overview + positioning + 联邦 naming) ============
body_kw = dict(align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体')
summary_block = [
    ("摘要", dict(style=H1, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=None, sa=Pt(10), font=None)),
    ("联邦智研（AISci）是一套面向科研全过程的 AI 科研生产力平台与科研智能体基础设施，而非单一的对话式“科研助手”。系统以国产大模型 Qwen 为底座，将“矛盾→拆解→知识空白→假设→验证”的科学逻辑链封装为可复现、可审计、可人机协同的自动化 Pipeline，帮助科研团队把“提出问题—检索证据—生成假设—验证假设—形成报告”的冗长流程压缩为小时级的可追溯工作流。", body_kw),
    ("关于名称中的“联邦”：本项目指的不是把模型参数分散到各机构的“联邦学习”本身，而是“多源证据的联邦（Federation）”与“跨团队 / 跨机构科研协作”——即把分散在 arXiv、开放数据集、用户私有文献库中的异构证据，在统一的质量门禁与证据链治理下“联邦”为一条可证伪、可溯源的科学假设。联邦学习（FL）只是我们验证系统通用性的一个演示场景（见第六章案例三），并非核心技术的全部。", body_kw),
    ("系统的技术内核是四项方法创新：① 多轮证据链迭代推理引擎（以 Fact 白名单从机制上杜绝幻觉引用）；② 布尔质量门禁（以 0/1 硬门禁替代连续评分，作为可审计质量闸）；③ 统一反馈中心 + 人在回路（HITL）的协同闭环；④ 反事实预演机制（L0 定性证伪，保护科研资源）。围绕这四项创新，系统构建了 6 个智能体、约 70 个可复用技能与 11 种阶段特定质量门禁，并在第七章给出结果展示与反馈迭代的完整量化。", body_kw),
    # ITEM 2 : 核心创新点 (centralized)
    ("核心创新点", dict(style=H1, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=None, sa=Pt(10), font=None)),
    ("创新一　多轮证据链迭代推理引擎（EvidenceChainBuilder）：构建“科学声明 ↔ 文献事实”的双向溯源与矛盾检测，以 Fact 白名单强制约束从机制层面杜绝幻觉引用，使每一条假设都可追溯、可证伪。", dict(align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体')),
    ("创新二　布尔质量门禁系统（Boolean Gate）：以 0/1 硬门禁替代常见的 0–100 连续评分，作为可审计的质量闸，在科研自动化评测中较为少见，构成方法层面的创新。", dict(align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体')),
    ("创新三　统一反馈中心 + 人在回路（Feedback Hub + HITL）：将某一阶段的发现经统一反馈中心传递给后续阶段，并在关键节点设置人工门控，实现“人机协同的闭环”，在自动化与可控性之间取得平衡。", dict(align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体')),
    ("创新四　反事实预演机制（Counterfactual Preview）：在假设评审与实验设计之间插入 L0 级定性证伪层，过滤不可证伪或缺乏决策影响的场景，阻断无价值实验，保护科研资源。", dict(align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体')),
]
block_before(ch1, summary_block)

# ============ ITEM 10 : in-body figure references ============
new_para(p21_end, "系统整体分层架构如图 2-1 所示，详细架构图见附录。",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体', sa=Pt(6))
new_para(p31_intro, "七阶段 Pipeline 的流程示意如图 3-1 所示，详细流程图见附录。",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体', sa=Pt(6))
new_para(p41_end, "证据链迭代推理引擎的流程如图 4-1 所示，详细流程图见附录。",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体', sa=Pt(6))

# ============ ITEM 7 : remove repeated Pipeline listing -> reference ============
p_pipeline_arrow.runs[0].text = ("七个阶段依次为：问题理解、文献挖掘、知识缺口发现、假设生成、假设评审、迭代实验、报告生成；"
                                 "各阶段的输入、产出与质量门禁机制详见第三章 3.1 节，本文不再赘述。")

# ============ ITEM 8 : bullet summaries after key subsections ============
bul = dict(align=WD_ALIGN_PARAGRAPH.LEFT, first_line=None, left=Pt(24), font='宋体', sa=Pt(3))
new_para(p12_end, "• 证据溯源：以 Fact 白名单约束，确保每条假设可追溯至真实文献，杜绝虚构引用。", **bul)
new_para(p12_end, "• 质量控制：以 11 种布尔质量门禁替代连续评分，在关键节点做离散化质量判定。", **bul)
new_para(p12_end, "• 迭代优化：以科学自迭代编排器实现“证据弱→补文献→重跑 Pipeline”的闭环精化。", **bul)

new_para(p41_end, "• 核心贡献：以 Fact 白名单 + 矛盾检测，从机制上约束 LLM 幻觉，使假设可证伪、可溯源。", **bul)
new_para(p42_end, "• 核心贡献：以 0/1 硬门禁替代连续评分，提供可审计、可停滞的质量决策边界。", **bul)
new_para(p43_end, "• 核心贡献：统一反馈中心跨阶段传递约束，HITL 门控在自动化与可控性间取得平衡。", **bul)
new_para(p44_end, "• 核心贡献：L0 定性证伪过滤无效反事实场景，阻断无价值实验，节约科研资源。", **bul)

new_para(p71_end, "• 小结：假设质量通过“领域对齐 → 新颖性审查 → 锦标赛排序 → 人工审核”四层机制保障，兼顾自动筛选与人工把关。", **bul)

# shorten two long sentences (ITEM 8)
para19.runs[0].text = ("针对上述挑战，本项目“联邦智研”（AISci）提出一种基于国产大模型 Qwen 的多智能体科研自动化系统，"
                       "以“矛盾→拆解→空白→假设→验证”的科学逻辑链为核心，构建从研究问题到可验证假设的全流程自动化 Pipeline，"
                       "并通过证据链迭代推理、布尔质量门禁、反事实预演等机制，系统性解决 LLM 在科研中的幻觉引用与质量控制问题。")
para85.runs[0].text = ("针对 LLM 生成的实验计划“看起来合理但实际无法执行”的问题，系统设计了实验计划可执行性 Gate："
                       "从实验设计的假设 / 方法 / 步骤中提取所需数据列信号，与可用数据列做精确 + 模糊匹配后按公式评分"
                       "（score = 40 + 25·has_data + 15·has_steps + 10·has_metrics + min(10, matched×2) − min(30, blockers×10)），"
                       "通过条件为 score≥60、无 blockers、有 steps、有 metrics。")

# ============ ITEM 6 : reduce Chapter 8 weight ============
ch8_head.runs[0].text = "八、应用示范：科学影响力预测系统"
new_para(p82_end, "科学影响力预测系统的四维评估模型、生命周期预测、偏差分析、组合评分、不确定量化及与主系统集成的技术细节，"
         "因篇幅所限移入附录二，本节约述设计原则与能力边界。",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=Pt(24), font='宋体', sa=Pt(6))

appx_head = new_para(appx_img4, "附录二：科学影响力预测系统技术细节",
                     style=H2, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=None, sa=Pt(8), font=None)
anchor = appx_head._p
for p in detail_paras:
    clone = copy.deepcopy(p._p)
    anchor.addnext(clone)
    anchor = clone
for p in detail_paras:
    p._p.getparent().remove(p._p)

doc.save(PATH)
print("EDITS APPLIED. Saved:", PATH)
