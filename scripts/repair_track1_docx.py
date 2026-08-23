# -*- coding: utf-8 -*-
"""Repair Track-1 template: replace corrupt settings/footnotes/endnotes."""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pythoncom
import win32com.client

BROKEN = Path(
    r"d:/Workplace/AISci/output/提交/模板/赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx"
)
DONOR = Path(r"d:/Workplace/AISci/output/提交/方向 1A1.docx")
TRACK2 = Path(
    r"d:/Workplace/AISci/output/提交/模板/赛道二-方向1B-科研影响力分析与偏差解释-提交要求及模板.docx"
)
COPY = Path(
    r"d:/Workplace/AISci/output/提交/赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板 副本.docx"
)

POISON = ("word/settings.xml", "word/footnotes.xml", "word/endnotes.xml")


def word_open(path: Path) -> str:
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(
            str(path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        n = doc.Paragraphs.Count
        doc.Close(False)
        return f"OK paras={n}"
    except Exception as e:
        detail = e.args[2][2] if len(getattr(e, "args", ())) > 2 else str(e)
        return f"FAIL {detail}"
    finally:
        word.Quit()
        pythoncom.CoUninitialize()


def repair(src: Path, donor: Path, dst: Path) -> None:
    with ZipFile(donor) as zd:
        replacements = {name: zd.read(name) for name in POISON}
    with ZipFile(src) as zin, ZipFile(dst, "w", ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = replacements[name] if name in replacements else zin.read(name)
            zout.writestr(name, data)


def diagnose_notes(path: Path) -> None:
    if not path.exists():
        print("missing", path)
        return
    with ZipFile(path) as z:
        settings = z.read("word/settings.xml").decode("utf-8", errors="replace")
        fn = z.read("word/footnotes.xml").decode("utf-8", errors="replace")
        en = z.read("word/endnotes.xml").decode("utf-8", errors="replace")
    import re

    print(path.name)
    print("  settings footnote ids", re.findall(r"<w:footnote w:id=\"(-?\d+)\"", settings))
    print("  settings endnote ids", re.findall(r"<w:endnote w:id=\"(-?\d+)\"", settings))
    print("  footnotes.xml ids", re.findall(r"w:id=\"(-?\d+)\"", fn))
    print("  endnotes.xml ids", re.findall(r"w:id=\"(-?\d+)\"", en))
    print("  word_open", word_open(path))


def main() -> None:
    print("=== before ===")
    diagnose_notes(BROKEN)
    diagnose_notes(TRACK2)
    diagnose_notes(COPY)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = BROKEN.with_name(BROKEN.stem + f".corrupt_bak_{stamp}.docx")
    shutil.copy2(BROKEN, bak)
    print("backup", bak.name)

    tmp = BROKEN.with_name(BROKEN.stem + "._repaired.docx")
    repair(BROKEN, DONOR, tmp)
    print("=== repaired tmp ===")
    print(tmp.name, word_open(tmp))

    # verify content preserved
    from docx import Document

    old = Document(str(BROKEN))
    new = Document(str(tmp))
    print("paras", len(old.paragraphs), "->", len(new.paragraphs))
    print("inline", len(old.inline_shapes), "->", len(new.inline_shapes))
    # spot-check a filled paragraph
    for p in new.paragraphs:
        if "四步" in (p.text or ""):
            print("spot", (p.text or "")[:80])
            break

    shutil.move(str(tmp), str(BROKEN))
    print("=== after replace ===")
    print(BROKEN.name, word_open(BROKEN))

    if COPY.exists():
        bak2 = COPY.with_name(COPY.stem + f".corrupt_bak_{stamp}.docx")
        shutil.copy2(COPY, bak2)
        repair(COPY, DONOR, COPY.with_name(COPY.stem + "._repaired.docx"))
        shutil.move(str(COPY.with_name(COPY.stem + "._repaired.docx")), str(COPY))
        print("copy repaired", word_open(COPY))


if __name__ == "__main__":
    main()
