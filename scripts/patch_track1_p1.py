# -*- coding: utf-8 -*-
"""Finish P1: 300-word intro + Qwen tech note, aligned to 1A method chain."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx import Document
from retune_track1_1a_theme import (  # noqa: E402
    HZ,
    PARAS,
    TPL,
    set_cell,
    set_fill,
    zh_len,
)

INTRO_300 = (
    "联邦智研以国产大模型 Qwen（阿里云百炼）为底座，证明方向 1A 需要的不是答对 125 题，"
    "而是四步方法链：科学问题如何变成可处理知识缺口；证据（含反对证据）如何约束假设；"
    "候选如何比较筛选并写成可检验计划；第一轮真实结果如何原样保留后再进入第二轮。"
    "七阶段依次完成问题理解、文献挖掘、知识缺口、假设生成、假设评审、迭代实验与报告生成。"
    "Fact 白名单与反证检索禁止无依据生成，门禁与锦标赛留下筛选理由，可执行性检查挡住空泛计划。"
    "125 题用于检验该方法链能否跨学科跑通；代表案例 sjtu_q_087 保留第一轮不可执行计划，不用事后美化版代替。"
)

QWEN_300 = (
    "系统以 Qwen 为唯一基座，经阿里云百炼调用 qwen-3.7max、qwen-3.6max、qwen-3.7plus、qwen-3.6plus，"
    "分别承担问题拆解、事实与反证抽取、缺口识别、假设生成、比较评审和计划撰写。"
    "提供给模型的不是「请直接回答这道科学问题」，而是按方法链打包的上下文："
    "科学问题、已有证据、反对证据、关键约束、历史结果与人工反馈；引用只允许白名单 fact_id。"
    "因此 Qwen 产出的是可处理缺口、带反对证据的多候选、筛选理由和可检验计划，而不是口号题的是/否答案。"
    "第一轮上下文与输出写入审计链后冻结，第二轮只追加新证据与反馈，不覆盖原件。"
)

TOPIC = "赛道一-方向1A-科学假设生成与研究计划设计（主证明）；方向1B科研影响力分析为延伸，不作为本方向主证据"


def _must_le_300(name: str, text: str) -> None:
    n = zh_len(text)
    print(f"{name} {n}/300")
    if n > 300:
        raise SystemExit(f"{name} 超出 300 字: {n}")


def patch_huizong_p1(doc: Document) -> int:
    mapping = {
        "R3 参赛作品简介": INTRO_300,
        "参赛作品简介（300字内）：": INTRO_300,
        "R4 Qwen及AI技术说明": QWEN_300,
        "Qwen及AI技术说明（300字内）：": QWEN_300,
        "R2 参赛选题：": TOPIC,
        "核心主张·具体问题：": PARAS[10].split("：", 1)[1],
        "核心主张·核心方法：": PARAS[11].split("：", 1)[1],
        "核心主张·结果1：": PARAS[12].split("：", 1)[1],
        "核心主张·结果2：": PARAS[13].split("：", 1)[1],
        "核心主张·结果3：": PARAS[14].split("：", 1)[1],
        "核心主张·主要局限：": PARAS[15].split("：", 1)[1],
    }
    n = 0
    for p in doc.paragraphs:
        raw = "".join((r.text or "") for r in p.runs)
        if raw.startswith("===="):
            continue
        for label, answer in mapping.items():
            if raw.startswith(label) or (label in raw[:40]):
                # keep the original prefix up to first colon of this field
                if "：" in raw[: len(label) + 5]:
                    head = raw.split("：", 1)[0] + "："
                else:
                    head = label if label.endswith("：") else label + "："
                set_fill(p, head + answer)
                n += 1
                break
    return n


def main() -> None:
    _must_le_300("简介", INTRO_300)
    _must_le_300("Qwen说明", QWEN_300)

    tpl = Document(str(TPL))
    set_cell(tpl.tables[0].rows[2].cells[1], TOPIC)
    set_cell(tpl.tables[0].rows[3].cells[1], INTRO_300)
    set_cell(tpl.tables[0].rows[4].cells[1], QWEN_300)
    tpl.save(str(TPL))
    print("saved template P1 table")

    hz = Document(str(HZ))
    n = patch_huizong_p1(hz)
    hz.save(str(HZ))
    print("huizong p1", n)


if __name__ == "__main__":
    main()
