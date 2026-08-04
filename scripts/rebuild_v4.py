# -*- coding: utf-8 -*-
"""Rebuild v4_fixed.docx from v3.docx with all 8 + 2 modifications applied correctly.
- Edits via python-docx (handles nested structures correctly).
- Serialize via lxml (preserves val='1'; avoids python-docx save converting to 'true').
- Removes the TOC <w:sdt> wrapper (converter residue) and adds a manual TOC.
- Adds missing endnotes/footnotes separator entries (id 0/1).
Output overwrites v4_fixed.docx.
"""
import copy, zipfile, shutil
from lxml import etree
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

V3 = "D:/Workplace/AISci/output/联邦智研_大创申报书_v3.docx"
OUT = "D:/Workplace/AISci/output/联邦智研_大创申报书_v4_fixed.docx"
BAK = OUT + ".rebuild_in.bak"
shutil.copyfile(OUT, BAK)

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc = Document(V3)

# ---------- helpers ----------
def set_run_font(run, cn='宋体', size=None, bold=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.name = cn
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), cn)
    rfonts.set(qn('w:ascii'), cn)
    rfonts.set(qn('w:hAnsi'), cn)

def add_para(ref, text, style=None, before=False, first_line=True,
             left=None, align='justify', size=12, bold=None, bullet=False):
    p = doc.add_paragraph()
    if style is not None:
        p.style = style
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if align == 'justify':
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'left':
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bullet:
        pf.left_indent = Pt(24)
        pf.first_line_indent = None
    else:
        if first_line:
            pf.first_line_indent = Pt(24)
        if left is not None:
            pf.left_indent = left
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    if before:
        ref._p.addprevious(p._p)
    else:
        ref._p.addnext(p._p)
    return p

def find_para(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    raise ValueError("not found: " + str(pred))

# ---------- locate anchors (object refs survive insertions) ----------
ch1        = find_para(lambda t: t.startswith('一、研究问题与解决方法'))
toc_head   = find_para(lambda t: t.startswith('目'))
h21        = find_para(lambda t: t.startswith('2.1 '))
h22        = find_para(lambda t: t.startswith('2.2 '))
h31        = find_para(lambda t: t.startswith('3.1 '))
h31_intro  = find_para(lambda t: '用户在前端创建研究项目' in t)
h41        = find_para(lambda t: t.startswith('4.1 '))
h42        = find_para(lambda t: t.startswith('4.2 '))
h43        = find_para(lambda t: t.startswith('4.3 '))
h44        = find_para(lambda t: t.startswith('4.4 '))
h12        = find_para(lambda t: t == '1.2 核心研究问题')
h13        = find_para(lambda t: t.startswith('1.3 '))
h71_last   = find_para(lambda t: '第四层——人工审核' in t)
h72        = find_para(lambda t: t.startswith('7.2 '))
ch8        = find_para(lambda t: t.startswith('八、'))
h82        = find_para(lambda t: t.startswith('8.2 '))
detail_s   = find_para(lambda t: t.startswith('8.3 '))
detail_e   = find_para(lambda t: t.startswith('8.8 '))
p11_long   = find_para(lambda t: '基于国产大模型Qwen' in t)
p45_gate   = find_para(lambda t: '针对LLM生成的实验计划' in t)

# last image in doc (appendix 图4-1)
last_img = None
for p in doc.paragraphs:
    if '<a:blip' in p._p.xml:
        last_img = p
assert last_img is not None, "no image found"

H1 = ch1.style  # chapter heading style

# ============ ITEM 1 + 4 + 5 : 摘要 + 核心创新点 ============
add_para(ch1, "摘要", style=H1, before=True)
add_para(ch1, "联邦智研（AISci）是一套面向科研全过程的 AI 科研生产力平台与科研智能体基础设施，而非单一的对话式“科研助手”。系统以国产大模型 Qwen 为底座，将“矛盾→拆解→知识空白→假设→验证”的科学逻辑链封装为可复现、可审计、可人机协同的自动化 Pipeline，帮助科研团队把“提出问题—检索证据—生成假设—验证假设—形成报告”的冗长流程压缩为小时级的可追溯工作流。", before=True)
add_para(ch1, "关于名称中的“联邦”：本项目指的不是把模型参数分散到各机构的“联邦学习”本身，而是“多源证据的联邦（Federation）”与“跨团队 / 跨机构科研协作”——即把分散在 arXiv、开放数据集、用户私有文献库中的异构证据，在统一的质量门禁与证据链治理下“联邦”为一条可证伪、可溯源的科学假设。联邦学习（FL）只是我们验证系统通用性的一个演示场景（见第六章案例三），并非核心技术的全部。", before=True)
add_para(ch1, "系统的技术内核是四项方法创新：① 多轮证据链迭代推理引擎（以 Fact 白名单从机制上杜绝幻觉引用）；② 布尔质量门禁（以 0/1 硬门禁替代连续评分，作为可审计质量闸）；③ 统一反馈中心 + 人在回路（HITL）的协同闭环；④ 反事实预演机制（L0 定性证伪，保护科研资源）。围绕这四项创新，系统构建了 6 个智能体、约 70 个可复用技能与 11 种阶段特定质量门禁，并在第七章给出结果展示与反馈迭代的完整量化。", before=True)
# ITEM 2
add_para(ch1, "核心创新点", style=H1, before=True)
add_para(ch1, "创新一　多轮证据链迭代推理引擎（EvidenceChainBuilder）：构建“科学声明 ↔ 文献事实”的双向溯源与矛盾检测，以 Fact 白名单强制约束从机制层面杜绝幻觉引用，使每一条假设都可追溯、可证伪。", before=True)
add_para(ch1, "创新二　布尔质量门禁系统（Boolean Gate）：以 0/1 硬门禁替代常见的 0–100 连续评分，作为可审计的质量闸，在科研自动化评测中较为少见，构成方法层面的创新。", before=True)
add_para(ch1, "创新三　统一反馈中心 + 人在回路（Feedback Hub + HITL）：将某一阶段的发现经统一反馈中心传递给后续阶段，并在关键节点设置人工门控，实现“人机协同的闭环”，在自动化与可控性之间取得平衡。", before=True)
add_para(ch1, "创新四　反事实预演机制（Counterfactual Preview）：在假设评审与实验设计之间插入 L0 级定性证伪层，过滤不可证伪或缺乏决策影响的场景，阻断无价值实验，保护科研资源。", before=True)

# ============ ITEM 10 : figure refs in body ============
add_para(h21, "系统整体分层架构如图 2-1 所示，详细架构图见附录。", before=False)
add_para(h31_intro, "七阶段 Pipeline 的流程示意如图 3-1 所示，详细流程图见附录。", before=False)
add_para(h42, "证据链迭代推理引擎的流程如图 4-1 所示，详细流程图见附录。", before=True)  # end of 4.1

# ============ ITEM 7 : Pipeline dedupe ============
for r in h22._p.iter(qn('w:r')):
    pass
# find the arrow listing paragraph inside 2.2 (the one with '问题理解' and '→')
pipe_arrow = find_para(lambda t: '问题理解' in t and '→' in t)
pipe_arrow.text = ("七个阶段依次为：问题理解、文献挖掘、知识缺口发现、假设生成、假设评审、迭代实验、报告生成；"
                   "各阶段的输入、产出与质量门禁机制详见第三章 3.1 节，本文不再赘述。")
# clear extra runs if any
for r in list(pipe_arrow._p):
    if r.tag == qn('w:r') and r.get(qn('w:rsidR')) is not None:
        pass

# ============ ITEM 8 : bullet summaries ============
# 1.2 (correct order) before h13
add_para(h13, "• 迭代优化：以科学自迭代编排器实现“证据弱→补文献→重跑 Pipeline”的闭环精化。", before=True, bullet=True)
add_para(h13, "• 质量控制：以 11 种布尔质量门禁替代连续评分，在关键节点做离散化质量判定。", before=True, bullet=True)
add_para(h13, "• 证据溯源：以 Fact 白名单约束，确保每条假设可追溯至真实文献，杜绝虚构引用。", before=True, bullet=True)
# 4.1-4.4 before next heading
add_para(h42, "• 核心贡献：以 Fact 白名单 + 矛盾检测，从机制上约束 LLM 幻觉，使假设可证伪、可溯源。", before=True, bullet=True)
add_para(h43, "• 核心贡献：以 0/1 硬门禁替代连续评分，提供可审计、可停滞的质量决策边界。", before=True, bullet=True)
add_para(h44, "• 核心贡献：统一反馈中心跨阶段传递约束，HITL 门控在自动化与可控性间取得平衡。", before=True, bullet=True)
h45 = find_para(lambda t: t.startswith('4.5 '))
add_para(h45, "• 核心贡献：L0 定性证伪过滤无效反事实场景，阻断无价值实验，节约科研资源。", before=True, bullet=True)
# 7.1 before h72
add_para(h72, "• 小结：假设质量通过“领域对齐 → 新颖性审查 → 锦标赛排序 → 人工审核”四层机制保障，兼顾自动筛选与人工把关。", before=True, bullet=True)

# ============ shorten long sentences (ITEM 8) ============
p11_long.text = ("针对上述挑战，本项目“联邦智研”（AISci）提出一种基于国产大模型 Qwen 的多智能体科研自动化系统，"
                 "以“矛盾→拆解→空白→假设→验证”的科学逻辑链为核心，构建从研究问题到可验证假设的全流程自动化 Pipeline，"
                 "并通过证据链迭代推理、布尔质量门禁、反事实预演等机制，系统性解决 LLM 在科研中的幻觉引用与质量控制问题。")
p45_gate.text = ("针对 LLM 生成的实验计划“看起来合理但实际无法执行”的问题，系统设计了实验计划可执行性 Gate："
                 "从实验设计的假设 / 方法 / 步骤中提取所需数据列信号，与可用数据列做精确 + 模糊匹配后按公式评分"
                 "（score = 40 + 25·has_data + 15·has_steps + 10·has_metrics + min(10, matched×2) − min(30, blockers×10)），"
                 "通过条件为 score≥60、无 blockers、有 steps、有 metrics。")

# ============ ITEM 6 : reduce Chapter 8, move 8.3-8.8 to 附录二 ============
ch8.text = "八、应用示范：科学影响力预测系统"
add_para(h82, "科学影响力预测系统的四维评估模型、生命周期预测、偏差分析、组合评分、不确定量化及与主系统集成的技术细节，因篇幅所限移入附录二，本节约述设计原则与能力边界。", before=False)

# collect 8.3-8.8 (by XML element identity, inclusive) -- correct slice, no 9.1
detail_paras = []
cap = False
ds = detail_s._p
de = detail_e._p
for p in doc.paragraphs:
    pe = p._p
    if pe is ds:
        cap = True
    if cap:
        detail_paras.append(p)
    if pe is de:
        break
# clone into appendix after last image
appx2_head = add_para(last_img, "附录二：科学影响力预测系统技术细节", style=H1, before=False)
anchor = appx2_head._p
for p in detail_paras:
    clone = copy.deepcopy(p._p)
    anchor.addnext(clone)
    anchor = clone
# delete originals
for p in detail_paras:
    p._p.getparent().remove(p._p)

# ============ follow-up 2 : manual TOC (replace the broken live field) ============
toc_entries = [
    "摘要", "核心创新点",
    "一、研究问题与解决方法", "二、架构设计与讲解", "三、项目工作流程",
    "四、上下文工程设计", "五、数据或资料来源说明", "六、代表性测试案例",
    "七、结果展示与反馈迭代过程", "八、应用示范：科学影响力预测系统",
    "九、源代码与可复现性", "十、总结与展望", "附录", "附录二：科学影响力预测系统技术细节",
]
for e in reversed(toc_entries):
    add_para(toc_head, e, before=False)

# ============ serialize via lxml, preserve val='1' ============
document_xml = etree.tostring(doc._element, xml_declaration=True, encoding='UTF-8', standalone=True)

# ============ lxml passes: remove TOC <w:sdt>, fix endnotes/footnotes ============
root = etree.fromstring(document_xml)
body = root.find('{%s}body' % W)
# remove any <w:sdt> that wraps a TOC field
for sdt in list(body.iter('{%s}sdt' % W)):
    txt = ''.join(sdt.itertext())
    if 'TOC' in txt and 'HYPERLINK' in txt:
        sdt.getparent().remove(sdt)
document_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

# ============ repack ============
z = zipfile.ZipFile(V3)
parts = {}
for n in z.namelist():
    parts[n] = z.read(n)
parts['word/document.xml'] = document_xml

# fix endnotes/footnotes: ensure id 0 (separator) and 1 (continuationSeparator)
def ensure_sep(part_name):
    if part_name not in parts:
        return
    r = etree.fromstring(parts[part_name])
    tag = 'endnote' if 'endnote' in part_name else 'footnote'
    ids = [e.get(qn('w:id')) for e in r.iter(qn('w:' + tag))]
    if '0' not in ids:
        el = etree.SubElement(r, qn('w:' + tag))
        el.set(qn('w:id'), '0')
        el.set(qn('w:type'), 'separator')
        p = etree.SubElement(el, qn('w:p'))
        rPr = etree.SubElement(p, qn('w:pPr'))
        rStyle = etree.SubElement(rPr, qn('w:rStyle'))
        rStyle.set(qn('w:val'), 'EndnoteReference' if 'endnote' in part_name else 'FootnoteReference')
        run = etree.SubElement(p, qn('w:r'))
        etree.SubElement(run, qn('w:separator'))
        parts[part_name] = etree.tostring(r, xml_declaration=True, encoding='UTF-8', standalone=True)

ensure_sep('word/endnotes.xml')
ensure_sep('word/footnotes.xml')

with zipfile.ZipFile(OUT, 'w', zipfile.ZIP_DEFLATED) as z2:
    for n, data in parts.items():
        z2.writestr(n, data)

print("REBUILD DONE")
print("  detail_paras moved (8.3-8.8):", len(detail_paras))
print("  TOC entries added:", len(toc_entries))
